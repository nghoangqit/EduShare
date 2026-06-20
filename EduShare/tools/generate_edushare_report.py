from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Bao_cao_DACS2_EduShare_sequence_hinh_anh.docx"


def set_run(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def paragraph(doc, text="", style=None, align=None, bold=False, italic=False, size=13):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    set_run(run, size=15 if level == 1 else 14, bold=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run(run)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run(run)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10.5)
    return p


def image_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"[[CẦN THÊM HÌNH ẢNH: {text}]]")
    set_run(run, size=12, bold=True, italic=True)
    run.font.color.rgb = RGBColor(192, 0, 0)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run(r, size=12)
    doc.add_paragraph()
    return table


def page_break(doc):
    doc.add_page_break()


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(13)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.color.rgb = RGBColor(0, 0, 0)


def cover(doc):
    paragraph(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN &", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    paragraph(doc, "TRUYỀN THÔNG VIỆT - HÀN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    paragraph(doc, "Khoa Khoa Học Máy Tính", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    paragraph(doc, "ĐỒ ÁN CƠ SỞ 2", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    doc.add_paragraph()
    paragraph(
        doc,
        "XÂY DỰNG ỨNG DỤNG EDU SHARE HỖ TRỢ MUA BÁN SÁCH VÀ DỤNG CỤ HỌC TẬP",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=17,
    )
    doc.add_paragraph()
    doc.add_paragraph()
    paragraph(doc, "Sinh viên thực hiện:      ........................................", size=13)
    paragraph(doc, "Lớp:                      ........................................", size=13)
    paragraph(doc, "MSSV:                     ........................................", size=13)
    paragraph(doc, "Giảng viên hướng dẫn:     ........................................", size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    paragraph(doc, "Đà Nẵng, tháng 05 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    page_break(doc)


def acknowledgements(doc):
    heading(doc, "LỜI CẢM ƠN", 1)
    texts = [
        "Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô đã hướng dẫn, góp ý và tạo điều kiện cho em trong quá trình thực hiện đồ án.",
        "Em xin cảm ơn Khoa Khoa Học Máy Tính đã cung cấp nền tảng kiến thức về lập trình ứng dụng, cơ sở dữ liệu, phân tích thiết kế hệ thống và triển khai phần mềm.",
        "Đề tài EduShare được xây dựng nhằm giải quyết nhu cầu trao đổi, mua bán sách và dụng cụ học tập giữa sinh viên. Trong quá trình thực hiện, em đã có cơ hội vận dụng Flutter, Firebase, PayOS và các kỹ thuật quản lý trạng thái để tạo ra một sản phẩm có tính ứng dụng thực tế.",
        "Do thời gian và kinh nghiệm còn hạn chế, báo cáo cũng như sản phẩm không tránh khỏi thiếu sót. Em rất mong nhận được sự góp ý của quý thầy cô để hệ thống được hoàn thiện hơn.",
        "Em xin trân trọng cảm ơn!",
    ]
    for t in texts:
        paragraph(doc, t)
    doc.add_paragraph()
    paragraph(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT)
    paragraph(doc, "........................................", align=WD_ALIGN_PARAGRAPH.RIGHT)
    page_break(doc)


def toc(doc):
    heading(doc, "MỤC LỤC", 1)
    items = [
        "Chương 1. Giới thiệu đề tài",
        "Chương 2. Cơ sở lý thuyết",
        "Chương 3. Phân tích và thiết kế hệ thống",
        "Chương 4. Thiết kế dữ liệu",
        "Chương 5. Thiết kế giao diện",
        "Chương 6. Kết luận",
        "Chương 7. Kiểm thử và đánh giá",
        "Phụ lục. Mô tả một số đoạn code quan trọng",
        "Chương 8. Tài liệu tham khảo",
    ]
    for item in items:
        paragraph(doc, item)
    page_break(doc)

    heading(doc, "DANH MỤC CÁC TỪ VIẾT TẮT", 1)
    add_table(
        doc,
        ["Từ viết tắt", "Ý nghĩa"],
        [
            ["UI", "User Interface - Giao diện người dùng"],
            ["UX", "User Experience - Trải nghiệm người dùng"],
            ["API", "Application Programming Interface"],
            ["CRUD", "Create, Read, Update, Delete"],
            ["QR", "Quick Response Code"],
            ["SDK", "Software Development Kit"],
            ["BaaS", "Backend as a Service"],
        ],
    )


def chapter_1(doc):
    heading(doc, "Chương 1. GIỚI THIỆU ĐỀ TÀI", 1)
    heading(doc, "1.1 Tên đề tài", 2)
    paragraph(doc, "Xây dựng ứng dụng EduShare hỗ trợ mua bán sách và dụng cụ học tập cho sinh viên.")

    heading(doc, "1.2 Lý do chọn đề tài", 2)
    for t in [
        "Trong môi trường đại học, sinh viên thường có nhu cầu mua lại sách, giáo trình, máy tính, dụng cụ học tập hoặc thanh lý các tài liệu không còn sử dụng.",
        "Các hình thức mua bán hiện tại chủ yếu diễn ra qua mạng xã hội, nhóm chat hoặc trao đổi trực tiếp. Những hình thức này còn hạn chế về tìm kiếm, quản lý đơn hàng, theo dõi người bán và thanh toán.",
        "EduShare được xây dựng nhằm tạo ra một nền tảng tập trung, giúp sinh viên đăng bán, tìm kiếm, đặt mua và thanh toán sản phẩm học tập một cách thuận tiện hơn.",
    ]:
        bullet(doc, t)

    heading(doc, "1.3 Mô tả sản phẩm", 2)
    paragraph(
        doc,
        "EduShare là ứng dụng di động được phát triển bằng Flutter, sử dụng Firebase làm nền tảng backend và tích hợp PayOS cho chức năng nạp tiền vào ví. Ứng dụng cho phép người dùng đăng nhập, quản lý hồ sơ, đăng bán sản phẩm, tìm kiếm sản phẩm, thêm vào giỏ hàng, thanh toán bằng ví EduShare, chat hỗ trợ và nhận thông báo.",
    )

    heading(doc, "1.4 Chức năng chính", 2)
    for t in [
        "Đăng ký, đăng nhập và quản lý hồ sơ cá nhân.",
        "Đăng bán sách, giáo trình, dụng cụ học tập.",
        "Tìm kiếm, xem chi tiết, yêu thích sản phẩm.",
        "Giỏ hàng, đặt hàng và thanh toán.",
        "Ví EduShare: nạp tiền, rút tiền, thanh toán bằng ví.",
        "Tích hợp PayOS để tạo mã thanh toán và tự động xác nhận nạp ví.",
        "Chat hỗ trợ với admin và phản hồi tự động.",
        "Thông báo trạng thái đơn hàng, ví và tin nhắn.",
        "Trang quản trị cho admin quản lý người dùng, sản phẩm, đơn hàng và ví.",
    ]:
        bullet(doc, t)

    heading(doc, "1.5 Công nghệ sử dụng", 2)
    add_table(
        doc,
        ["Công nghệ", "Vai trò trong hệ thống"],
        [
            ["Flutter", "Xây dựng giao diện ứng dụng di động đa nền tảng"],
            ["Dart", "Ngôn ngữ lập trình chính"],
            ["Firebase Authentication", "Xác thực người dùng"],
            ["Cloud Firestore", "Lưu trữ dữ liệu sản phẩm, đơn hàng, ví, chat, thông báo"],
            ["Firebase Messaging", "Hỗ trợ thông báo"],
            ["Provider", "Quản lý trạng thái đăng nhập và giỏ hàng"],
            ["PayOS", "Tạo mã thanh toán và kiểm tra trạng thái nạp ví"],
            ["Flutter Map", "Chọn vị trí giao hàng"],
        ],
    )

    heading(doc, "1.6 Phạm vi đề tài", 2)
    paragraph(
        doc,
        "Trong phạm vi đồ án, EduShare tập trung vào nhóm chức năng cốt lõi của một ứng dụng thương mại học đường: quản lý tài khoản, quản lý sản phẩm, giỏ hàng, đơn hàng, ví điện tử nội bộ, chat hỗ trợ và trang quản trị. Hệ thống được triển khai ở mức ứng dụng di động, dữ liệu lưu trữ trên Firebase, chưa triển khai backend riêng cho toàn bộ nghiệp vụ.",
    )
    bullet(doc, "Phạm vi người dùng: sinh viên, người bán và admin quản trị hệ thống.")
    bullet(doc, "Phạm vi sản phẩm: sách, giáo trình, máy tính, dụng cụ học tập và các tài liệu phục vụ học tập.")
    bullet(doc, "Phạm vi thanh toán: thanh toán khi nhận hàng và thanh toán bằng ví EduShare; ví có thể được nạp thông qua PayOS.")
    bullet(doc, "Phạm vi quản trị: quản lý người dùng, sản phẩm, đơn hàng, yêu cầu ví và thông báo.")

    heading(doc, "1.7 Đối tượng sử dụng", 2)
    add_table(
        doc,
        ["Đối tượng", "Nhu cầu sử dụng"],
        [
            ["Sinh viên mua hàng", "Tìm sách, dụng cụ học tập giá phù hợp; đặt hàng; thanh toán; theo dõi đơn hàng."],
            ["Sinh viên bán hàng", "Đăng bán sản phẩm đã qua sử dụng; quản lý sản phẩm; nhận tiền giải ngân."],
            ["Admin", "Kiểm duyệt dữ liệu, hỗ trợ người dùng, xử lý ví, quản lý đơn hàng và đảm bảo hệ thống hoạt động ổn định."],
        ],
    )

    heading(doc, "1.8 Ý nghĩa thực tiễn", 2)
    paragraph(
        doc,
        "EduShare có ý nghĩa thực tiễn trong môi trường sinh viên vì giúp tái sử dụng tài liệu học tập, giảm lãng phí, tiết kiệm chi phí và tạo kênh kết nối tập trung giữa người mua và người bán. Thay vì tìm kiếm rời rạc trên các nhóm mạng xã hội, người dùng có thể xem thông tin sản phẩm, giá, số lượng, người bán và trạng thái đơn hàng ngay trong một ứng dụng.",
    )
    paragraph(
        doc,
        "Về mặt kỹ thuật, đề tài giúp sinh viên rèn luyện khả năng xây dựng ứng dụng Flutter hoàn chỉnh, kết nối Firebase, tổ chức dữ liệu NoSQL, xử lý trạng thái, tích hợp thanh toán và thiết kế giao diện có trải nghiệm người dùng tốt.",
    )


def chapter_2(doc):
    heading(doc, "Chương 2. CƠ SỞ LÝ THUYẾT", 1)
    sections = [
        ("2.1 Flutter và Dart", "Flutter là framework mã nguồn mở do Google phát triển, cho phép xây dựng ứng dụng đa nền tảng từ một codebase. Dart là ngôn ngữ lập trình được Flutter sử dụng, hỗ trợ lập trình hướng đối tượng, bất đồng bộ và tối ưu cho giao diện phản hồi nhanh."),
        ("2.2 Firebase Authentication", "Firebase Authentication cung cấp cơ chế đăng nhập, quản lý phiên người dùng và xác thực tài khoản. Trong EduShare, Firebase Auth được dùng để xác định người dùng hiện tại và phân quyền thao tác dữ liệu."),
        ("2.3 Cloud Firestore", "Cloud Firestore là cơ sở dữ liệu NoSQL thời gian thực. Dữ liệu được tổ chức theo collection và document, phù hợp với các đối tượng như users, products, orders, walletRequests, conversations và notifications."),
        ("2.4 Provider", "Provider là thư viện quản lý trạng thái phổ biến trong Flutter. EduShare sử dụng Provider để quản lý trạng thái đăng nhập và giỏ hàng, giúp giao diện tự cập nhật khi dữ liệu thay đổi."),
        ("2.5 PayOS", "PayOS là nền tảng thanh toán hỗ trợ tạo link thanh toán, mã QR và kiểm tra trạng thái giao dịch. EduShare dùng PayOS cho chức năng nạp tiền vào ví."),
        ("2.6 UI/UX trong ứng dụng di động", "Ứng dụng di động cần giao diện rõ ràng, thao tác nhanh và phản hồi trực quan. EduShare sử dụng thẻ thông tin, thanh điều hướng, hiệu ứng chuyển tab và hiệu ứng thanh toán thành công để nâng cao trải nghiệm người dùng."),
    ]
    for title, body in sections:
        heading(doc, title, 2)
        paragraph(doc, body)

    heading(doc, "2.7 Kiến trúc dịch vụ trong ứng dụng", 2)
    paragraph(
        doc,
        "EduShare tách phần xử lý nghiệp vụ thành các service để tránh đặt quá nhiều logic trong giao diện. Cách tổ chức này giúp code dễ bảo trì và dễ mở rộng khi thêm chức năng mới.",
    )
    add_table(
        doc,
        ["Service", "Vai trò"],
        [
            ["FirebaseDataService", "Lớp trung tâm xử lý đọc/ghi dữ liệu Firestore, đơn hàng, ví, chat, thông báo."],
            ["PayosService", "Tạo link thanh toán PayOS, ký dữ liệu và kiểm tra trạng thái thanh toán."],
            ["NotificationSystemService", "Theo dõi thông báo của người dùng và hiển thị thông báo trong ứng dụng."],
            ["SupportBotService", "Sinh phản hồi hỗ trợ tự động cho một số câu hỏi phổ biến."],
        ],
    )

    heading(doc, "2.8 Bảo mật trong Firebase", 2)
    paragraph(
        doc,
        "Firebase Authentication chỉ xác thực danh tính người dùng, còn quyền đọc/ghi dữ liệu cần được giới hạn bằng Firestore Security Rules. Với các collection nhạy cảm như users, orders, walletRequests, hệ thống cần kiểm tra request.auth.uid để đảm bảo người dùng chỉ thao tác trên dữ liệu của mình, admin mới có quyền xử lý dữ liệu toàn hệ thống.",
    )
    bullet(doc, "Người dùng chỉ được tạo đơn hàng có buyerUid trùng với tài khoản hiện tại.")
    bullet(doc, "Người dùng chỉ được xem yêu cầu ví của chính mình.")
    bullet(doc, "Admin có quyền xem và xử lý yêu cầu ví, đơn hàng và sản phẩm.")
    bullet(doc, "Các thao tác cộng tiền ví cần được kiểm soát chặt chẽ và nên xử lý ở backend/cloud function trong môi trường production.")

    heading(doc, "2.9 Thanh toán và đối soát", 2)
    paragraph(
        doc,
        "Thanh toán là phần có yêu cầu chính xác cao. Khi tích hợp PayOS, ứng dụng cần tạo mã thanh toán, lưu mã đơn, kiểm tra trạng thái và chỉ cộng tiền khi trạng thái giao dịch là PAID. Trong phiên bản hiện tại, ứng dụng có cơ chế polling trạng thái PayOS từ client để phục vụ demo. Khi triển khai thực tế, nên chuyển phần kiểm tra và webhook PayOS sang backend để bảo mật API key và tránh giả mạo.",
    )

    heading(doc, "2.10 Giao diện phản hồi và animation", 2)
    paragraph(
        doc,
        "Ứng dụng di động cần phản hồi rõ ràng với thao tác của người dùng. Các animation như chuyển tab, hiệu ứng xác nhận thanh toán, loading indicator và snackbar giúp người dùng biết hệ thống đang xử lý gì, đã thành công hay cần thao tác lại.",
    )


def chapter_3(doc):
    heading(doc, "Chương 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    heading(doc, "3.1 Mô hình hệ thống", 2)
    paragraph(doc, "Hệ thống EduShare được tổ chức theo mô hình client - cloud backend:")
    bullet(doc, "Presentation Layer: các màn hình trong thư mục lib/screens như HomeScreen, SearchScreen, CartScreen, ProfileScreen.")
    bullet(doc, "State Management Layer: các provider trong lib/providers như AuthProvider và CartProvider.")
    bullet(doc, "Service Layer: FirebaseDataService, PayosService, NotificationSystemService.")
    bullet(doc, "Data Layer: các collection trong Cloud Firestore và các model trong lib/models.")

    heading(doc, "3.1.1 Cấu trúc thư mục dự án", 2)
    add_table(
        doc,
        ["Thư mục/File", "Chức năng"],
        [
            ["lib/models", "Định nghĩa các lớp dữ liệu: Product, UserProfile, WalletRequest, PurchaseRecord..."],
            ["lib/screens", "Chứa giao diện các màn hình của ứng dụng."],
            ["lib/services", "Chứa các lớp service xử lý Firebase, PayOS, thông báo, bot hỗ trợ."],
            ["lib/providers", "Quản lý trạng thái đăng nhập và giỏ hàng."],
            ["lib/utils", "Chứa hằng số giao diện, cấu hình admin, helper định dạng dữ liệu."],
            ["assets/images", "Lưu trữ hình ảnh tĩnh dùng trong ứng dụng."],
        ],
    )

    heading(doc, "3.2 Sơ đồ Use-case tổng quan", 2)
    paragraph(doc, "Các actor chính của hệ thống gồm: Khách, Người dùng, Người bán và Admin.")
    add_table(
        doc,
        ["Actor", "Use-case chính"],
        [
            ["Khách", "Đăng ký, đăng nhập, xem sản phẩm"],
            ["Người dùng", "Tìm kiếm, yêu thích, thêm giỏ hàng, đặt hàng, thanh toán, chat hỗ trợ"],
            ["Người bán", "Đăng sản phẩm, quản lý sản phẩm đang bán, nhận giải ngân"],
            ["Admin", "Quản lý người dùng, sản phẩm, đơn hàng, yêu cầu ví và thông báo"],
        ],
    )

    heading(doc, "3.3 Danh sách use-case", 2)
    use_cases = [
        "Đăng nhập/đăng ký",
        "Quản lý hồ sơ cá nhân",
        "Đăng bán sản phẩm",
        "Tìm kiếm sản phẩm",
        "Yêu thích sản phẩm",
        "Thêm sản phẩm vào giỏ hàng",
        "Đặt hàng",
        "Thanh toán bằng ví EduShare",
        "Nạp tiền vào ví qua PayOS",
        "Rút tiền từ ví",
        "Chat hỗ trợ admin",
        "Nhận thông báo",
        "Quản trị hệ thống",
    ]
    for item in use_cases:
        bullet(doc, item)

    heading(doc, "3.4 Đặc tả một số use-case", 2)
    add_table(
        doc,
        ["Use-case", "Actor", "Luồng xử lý chính"],
        [
            ["Đăng nhập", "Người dùng", "Nhập email/mật khẩu -> Firebase Auth xác thực -> tạo/lấy hồ sơ người dùng"],
            ["Đăng bán sản phẩm", "Người bán", "Nhập thông tin sản phẩm -> tạo Product -> lưu vào collection products"],
            ["Thanh toán bằng ví", "Người dùng", "Kiểm tra số dư -> tạo order -> trừ ví -> cập nhật tồn kho"],
            ["Nạp ví PayOS", "Người dùng", "Nhập số tiền -> tạo walletRequest -> tạo link PayOS -> kiểm tra trạng thái -> cộng ví"],
            ["Duyệt yêu cầu ví", "Admin", "Xem walletRequests -> xác nhận nạp/rút -> cập nhật trạng thái và thông báo"],
        ],
    )

    heading(doc, "3.4.1 Use-case Đăng nhập", 2)
    paragraph(doc, "Actor: Người dùng đã có tài khoản.")
    paragraph(doc, "Mục đích: Cho phép người dùng truy cập các chức năng cá nhân như giỏ hàng, ví, đơn hàng, chat và hồ sơ.")
    paragraph(doc, "Luồng chính:")
    numbered(doc, "Người dùng mở ứng dụng và nhập email, mật khẩu.")
    numbered(doc, "AuthProvider gọi Firebase Authentication để xác thực.")
    numbered(doc, "Sau khi đăng nhập thành công, PostAuthGate kiểm tra hồ sơ người dùng.")
    numbered(doc, "FirebaseDataService.ensureUserProfile tạo hồ sơ mới nếu chưa tồn tại.")
    numbered(doc, "Ứng dụng chuyển vào AppShell.")
    paragraph(doc, "File liên quan: lib/providers/auth_provider.dart, lib/screens/post_auth_gate.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.2 Use-case Đăng bán sản phẩm", 2)
    paragraph(doc, "Actor: Người dùng đã đăng nhập.")
    paragraph(doc, "Mục đích: Cho phép sinh viên đăng bán sách hoặc dụng cụ học tập.")
    paragraph(doc, "Luồng chính:")
    numbered(doc, "Người dùng bấm nút dấu cộng ở AppShell.")
    numbered(doc, "Ứng dụng mở AddProductScreen.")
    numbered(doc, "Người dùng nhập tên sản phẩm, giá, số lượng, mô tả, hình ảnh và trường học.")
    numbered(doc, "Ứng dụng tạo đối tượng Product.")
    numbered(doc, "FirebaseDataService.insertProduct lưu sản phẩm vào collection products.")
    paragraph(doc, "File liên quan: lib/screens/add_product_screen.dart, lib/models/product.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.3 Use-case Thêm sản phẩm vào giỏ hàng", 2)
    paragraph(doc, "Actor: Người mua.")
    paragraph(doc, "Mục đích: Cho phép người mua gom nhiều sản phẩm trước khi đặt hàng.")
    numbered(doc, "Người dùng xem sản phẩm trên HomeScreen hoặc SearchScreen.")
    numbered(doc, "Người dùng bấm thêm vào giỏ.")
    numbered(doc, "CartProvider.addToCart cập nhật danh sách CartItem.")
    numbered(doc, "CartScreen hiển thị tổng số lượng và tổng tiền.")
    paragraph(doc, "File liên quan: lib/providers/cart_provider.dart, lib/screens/cart_screen.dart, lib/models/cart_item.dart.")

    heading(doc, "3.4.4 Use-case Đặt hàng thanh toán khi nhận hàng", 2)
    paragraph(doc, "Actor: Người mua.")
    paragraph(doc, "Mục đích: Tạo đơn hàng với trạng thái chờ thanh toán khi nhận hàng.")
    numbered(doc, "Người dùng vào CartScreen và bấm thanh toán.")
    numbered(doc, "Ứng dụng kiểm tra thông tin giao hàng.")
    numbered(doc, "Người dùng chọn phương thức Thanh toán khi nhận hàng.")
    numbered(doc, "CartScreen gọi _completeOrder với paymentMethod là cod.")
    numbered(doc, "FirebaseDataService.createOrdersFromCart tạo document trong orders và cập nhật tồn kho.")
    paragraph(doc, "File liên quan: lib/screens/cart_screen.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.5 Use-case Thanh toán bằng ví EduShare", 2)
    paragraph(doc, "Actor: Người mua.")
    paragraph(doc, "Mục đích: Thanh toán nhanh bằng số dư ví nội bộ.")
    numbered(doc, "Người dùng chọn Thanh toán bằng ví EduShare.")
    numbered(doc, "Ứng dụng lấy walletBalance từ UserProfile.")
    numbered(doc, "Nếu số dư không đủ, hệ thống hiển thị thông báo lỗi.")
    numbered(doc, "Nếu đủ, FirebaseDataService.createWalletPaidOrdersFromCart tạo đơn hàng và trừ ví.")
    numbered(doc, "Giỏ hàng được xóa và màn hình hiển thị thành công.")
    paragraph(doc, "File liên quan: lib/screens/cart_screen.dart, lib/models/user_profile.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.6 Use-case Nạp tiền vào ví bằng PayOS", 2)
    paragraph(doc, "Actor: Người dùng.")
    paragraph(doc, "Mục đích: Nạp tiền vào ví EduShare thông qua mã QR/thanh toán PayOS.")
    numbered(doc, "Người dùng mở ProfileScreen và chọn Nạp tiền.")
    numbered(doc, "Ứng dụng gọi requestWalletDeposit để tạo WalletRequest.")
    numbered(doc, "PayosService.createWalletTopupLink tạo link thanh toán.")
    numbered(doc, "Màn _WalletTopupScreen hiển thị QR và tự kiểm tra trạng thái.")
    numbered(doc, "Nếu PayOS trả về PAID, _completeWalletDeposit cộng tiền vào ví.")
    paragraph(doc, "File liên quan: lib/screens/profile_screen.dart, lib/services/payos_service.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.7 Use-case Rút tiền từ ví", 2)
    paragraph(doc, "Actor: Người dùng có số dư ví và thông tin ngân hàng.")
    numbered(doc, "Người dùng mở Hồ sơ và chọn Rút tiền.")
    numbered(doc, "Ứng dụng kiểm tra tài khoản ngân hàng và số dư ví.")
    numbered(doc, "FirebaseDataService.requestWalletWithdrawal tạo yêu cầu rút và tạm trừ số dư.")
    numbered(doc, "Admin xử lý yêu cầu qua AdminDashboardScreen.")
    paragraph(doc, "File liên quan: lib/screens/profile_screen.dart, lib/screens/admin_dashboard_screen.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.4.8 Use-case Chat hỗ trợ", 2)
    numbered(doc, "Người dùng bấm Chat admin trong ProfileScreen.")
    numbered(doc, "FirebaseDataService.ensureAdminConversation tạo hoặc lấy cuộc trò chuyện.")
    numbered(doc, "ChatScreen hiển thị tin nhắn theo stream từ Firestore.")
    numbered(doc, "sendChatMessage lưu tin nhắn và tạo thông báo cho người nhận.")
    numbered(doc, "SupportBotService có thể phản hồi tự động cho câu hỏi phổ biến.")
    paragraph(doc, "File liên quan: lib/screens/chat_screen.dart, lib/screens/chat_list_screen.dart, lib/services/support_bot_service.dart.")

    heading(doc, "3.4.9 Use-case Quản trị hệ thống", 2)
    paragraph(doc, "Actor: Admin.")
    numbered(doc, "Admin đăng nhập bằng email được cấu hình trong AdminConfig.")
    numbered(doc, "ProfileScreen hiển thị lối vào AdminDashboardScreen.")
    numbered(doc, "Admin xem người dùng, sản phẩm, đơn hàng và yêu cầu ví.")
    numbered(doc, "Admin có thể khóa tài khoản, xóa sản phẩm, duyệt nạp/rút ví và giải ngân cho người bán.")
    paragraph(doc, "File liên quan: lib/screens/admin_dashboard_screen.dart, lib/utils/constants.dart, lib/services/firebase_data_service.dart.")

    heading(doc, "3.5 Phân tích các lớp chính", 2)
    add_table(
        doc,
        ["Lớp/Model", "Vai trò"],
        [
            ["UserProfile", "Lưu thông tin người dùng, ví, địa chỉ giao hàng và tài khoản ngân hàng"],
            ["Product", "Đại diện sản phẩm được đăng bán"],
            ["CartItem", "Đại diện một sản phẩm trong giỏ hàng"],
            ["PurchaseRecord", "Đại diện đơn hàng/lịch sử mua"],
            ["WalletRequest", "Đại diện yêu cầu nạp/rút ví"],
            ["AppNotification", "Đại diện thông báo trong hệ thống"],
            ["ChatConversation, ChatMessage", "Đại diện cuộc trò chuyện và tin nhắn"],
        ],
    )

    heading(doc, "3.6 Luồng xử lý chức năng nạp ví PayOS", 2)
    numbered(doc, "Người dùng mở Hồ sơ và chọn Nạp tiền.")
    numbered(doc, "Ứng dụng gọi requestWalletDeposit(requestedAmount) trong FirebaseDataService.")
    numbered(doc, "Hệ thống tạo document trong collection walletRequests với trạng thái pending.")
    numbered(doc, "PayosService.createWalletTopupLink tạo link/QR thanh toán PayOS.")
    numbered(doc, "Màn QR tự động kiểm tra trạng thái giao dịch theo chu kỳ.")
    numbered(doc, "Nếu PayOS trả về trạng thái PAID, hệ thống gọi _completeWalletDeposit.")
    numbered(doc, "Firestore transaction cập nhật walletRequest sang completed và cộng tiền vào walletBalance bằng FieldValue.increment.")

    heading(doc, "3.7 Luồng xử lý chức năng thanh toán bằng ví", 2)
    numbered(doc, "Người dùng chọn sản phẩm và thêm vào giỏ hàng.")
    numbered(doc, "Người dùng chọn thanh toán bằng ví EduShare.")
    numbered(doc, "CartScreen gọi createWalletPaidOrdersFromCart trong FirebaseDataService.")
    numbered(doc, "Hệ thống kiểm tra số dư ví và tồn kho sản phẩm.")
    numbered(doc, "Nếu hợp lệ, hệ thống tạo đơn hàng, trừ ví và cập nhật tồn kho.")
    numbered(doc, "Ứng dụng hiển thị thông báo/hiệu ứng đặt hàng thành công.")

    heading(doc, "3.8 Liên kết chức năng với mã nguồn", 2)
    add_table(
        doc,
        ["Chức năng", "File/Hàm chính"],
        [
            ["Đăng nhập", "AuthProvider, PostAuthGate, FirebaseDataService.ensureUserProfile"],
            ["Điều hướng tab", "AppShell._selectTab, AppShell._buildAnimatedTabBody"],
            ["Tìm kiếm", "FirebaseDataService.searchProducts"],
            ["Gợi ý sản phẩm", "FirebaseDataService.getRecommendedProducts"],
            ["Đăng sản phẩm", "AddProductScreen, FirebaseDataService.insertProduct"],
            ["Giỏ hàng", "CartProvider, CartScreen"],
            ["Tạo đơn hàng", "FirebaseDataService.createOrdersFromCart"],
            ["Thanh toán bằng ví", "FirebaseDataService.createWalletPaidOrdersFromCart"],
            ["Nạp ví", "ProfileScreen._showWalletDepositSheet, FirebaseDataService.requestWalletDeposit"],
            ["PayOS", "PayosService.createWalletTopupLink, PayosService.getPaymentLink"],
            ["Cộng ví", "FirebaseDataService._completeWalletDeposit"],
            ["Chat", "FirebaseDataService.sendChatMessage, ChatScreen"],
            ["Thông báo", "FirebaseDataService._createNotification, NotificationSystemService"],
            ["Admin", "AdminDashboardScreen và các hàm admin trong FirebaseDataService"],
        ],
    )

    heading(doc, "3.9 Mô tả các trạng thái quan trọng", 2)
    add_table(
        doc,
        ["Đối tượng", "Trạng thái", "Ý nghĩa"],
        [
            ["walletRequests", "pending", "Yêu cầu nạp/rút đang chờ xử lý"],
            ["walletRequests", "completed", "Yêu cầu ví đã hoàn tất"],
            ["walletRequests", "cancelled", "Người dùng đã hủy yêu cầu"],
            ["orders", "pending_cod", "Đơn hàng thanh toán khi nhận hàng"],
            ["orders", "awaiting_shipment", "Đơn đã thanh toán, chờ người bán giao"],
            ["orders", "delivered_pending_release", "Đơn đã giao, chờ admin giải ngân"],
            ["orders", "completed", "Đơn hàng hoàn tất"],
        ],
    )

    heading(doc, "3.10 Xử lý lỗi và phản hồi người dùng", 2)
    paragraph(doc, "Trong quá trình sử dụng, ứng dụng cần xử lý các tình huống như không đủ số dư ví, sản phẩm hết hàng, lỗi quyền Firestore, PayOS chưa xác nhận thanh toán hoặc mất kết nối mạng. EduShare sử dụng SnackBar, trạng thái loading và các kiểm tra điều kiện trước khi ghi dữ liệu để giảm lỗi thao tác.")
    bullet(doc, "Khi số dư ví không đủ, CartScreen hiển thị số tiền hiện có và số tiền cần thanh toán.")
    bullet(doc, "Khi sản phẩm hết hàng hoặc tồn kho thay đổi, hệ thống ném StateError và thông báo người dùng thử lại.")
    bullet(doc, "Khi PayOS chưa báo PAID, màn nạp ví tiếp tục kiểm tra và thông báo giao dịch chưa được xác nhận.")
    bullet(doc, "Khi thanh toán thành công, overlay thành công giúp người dùng nhận biết kết quả rõ ràng.")

    heading(doc, "3.11 Sơ đồ tuần tự (Sequence Diagram)", 2)
    paragraph(
        doc,
        "Phần này mô tả tuần tự tương tác giữa người dùng, giao diện Flutter, service trong ứng dụng, Firebase và PayOS đối với các chức năng quan trọng. Khi đưa vào báo cáo chính thức, có thể chuyển các sơ đồ dạng văn bản dưới đây thành hình UML bằng StarUML, draw.io hoặc PlantUML.",
    )

    heading(doc, "3.11.1 Sequence Diagram: Đăng nhập và khởi tạo hồ sơ", 2)
    code_block(
        doc,
        "User -> LoginScreen: Nhập email/mật khẩu\n"
        "LoginScreen -> AuthProvider: login(email, password)\n"
        "AuthProvider -> FirebaseAuth: signInWithEmailAndPassword()\n"
        "FirebaseAuth --> AuthProvider: UserCredential\n"
        "PostAuthGate -> FirebaseDataService: ensureUserProfile(user)\n"
        "FirebaseDataService -> Firestore.users: get(uid)\n"
        "alt Chưa có hồ sơ\n"
        "  FirebaseDataService -> Firestore.users: set(profile)\n"
        "end\n"
        "PostAuthGate -> AppShell: Hiển thị màn hình chính",
    )
    image_note(doc, "Vẽ lại sơ đồ Sequence Diagram Đăng nhập bằng công cụ UML và chèn tại đây.")

    heading(doc, "3.11.2 Sequence Diagram: Đăng bán sản phẩm", 2)
    code_block(
        doc,
        "Seller -> AppShell: Bấm nút +\n"
        "AppShell -> AddProductScreen: Mở form đăng sản phẩm\n"
        "Seller -> AddProductScreen: Nhập tên, giá, tồn kho, ảnh, mô tả\n"
        "AddProductScreen -> Product: Tạo đối tượng Product\n"
        "AddProductScreen -> FirebaseDataService: insertProduct(product)\n"
        "FirebaseDataService -> Firestore.products: set(product.toFirestore())\n"
        "Firestore.products --> FirebaseDataService: Ghi thành công\n"
        "AddProductScreen -> AppShell: Quay về trang chủ\n"
        "HomeScreen -> FirebaseDataService: getRecentProducts()\n"
        "FirebaseDataService --> HomeScreen: Danh sách sản phẩm mới",
    )
    image_note(doc, "Chèn ảnh màn hình form đăng sản phẩm và ảnh sản phẩm mẫu đã upload.")

    heading(doc, "3.11.3 Sequence Diagram: Thêm sản phẩm vào giỏ hàng", 2)
    code_block(
        doc,
        "User -> HomeScreen/SearchScreen: Chọn sản phẩm\n"
        "User -> ProductCard/ProductDetail: Bấm thêm vào giỏ\n"
        "ProductCard -> CartProvider: addToCart(product)\n"
        "CartProvider -> CartProvider: Cập nhật CartItem và tổng số lượng\n"
        "CartProvider --> AppShell: notifyListeners()\n"
        "AppShell -> BottomNavigation: Cập nhật badge giỏ hàng\n"
        "User -> CartScreen: Mở giỏ hàng\n"
        "CartScreen -> CartProvider: Lấy danh sách CartItem",
    )
    image_note(doc, "Chèn ảnh sản phẩm trong danh sách và ảnh giỏ hàng sau khi thêm sản phẩm.")

    heading(doc, "3.11.4 Sequence Diagram: Thanh toán bằng ví EduShare", 2)
    code_block(
        doc,
        "User -> CartScreen: Bấm thanh toán\n"
        "CartScreen -> FirebaseDataService: getCurrentUserProfile()\n"
        "FirebaseDataService --> CartScreen: UserProfile(walletBalance)\n"
        "CartScreen -> User: Hiển thị phương thức thanh toán\n"
        "User -> CartScreen: Chọn Thanh toán bằng ví EduShare\n"
        "CartScreen -> FirebaseDataService: createWalletPaidOrdersFromCart(items)\n"
        "FirebaseDataService -> Firestore.products: Kiểm tra tồn kho\n"
        "alt Số dư đủ và tồn kho đủ\n"
        "  FirebaseDataService -> Firestore.orders: Tạo đơn hàng\n"
        "  FirebaseDataService -> Firestore.users: Trừ walletBalance\n"
        "  FirebaseDataService -> Firestore.products: Cập nhật tồn kho\n"
        "  CartScreen -> CartProvider: clearCart()\n"
        "  CartScreen -> User: Hiển thị đặt hàng thành công\n"
        "else Không đủ điều kiện\n"
        "  CartScreen -> User: Hiển thị thông báo lỗi\n"
        "end",
    )
    image_note(doc, "Chèn ảnh màn hình chọn phương thức thanh toán chỉ còn COD và ví EduShare.")

    heading(doc, "3.11.5 Sequence Diagram: Nạp ví qua PayOS", 2)
    code_block(
        doc,
        "User -> ProfileScreen: Chọn Nạp tiền\n"
        "ProfileScreen -> FirebaseDataService: requestWalletDeposit(amount)\n"
        "FirebaseDataService -> Firestore.walletRequests: Tạo request pending\n"
        "FirebaseDataService -> PayosService: createWalletTopupLink(...)\n"
        "PayosService -> PayOS API: POST /v2/payment-requests\n"
        "PayOS API --> PayosService: checkoutUrl, qrCode, paymentLinkId\n"
        "FirebaseDataService -> Firestore.walletRequests: Cập nhật thông tin PayOS\n"
        "ProfileScreen -> _WalletTopupScreen: Hiển thị QR thanh toán\n"
        "loop Mỗi vài giây\n"
        "  _WalletTopupScreen -> FirebaseDataService: autoConfirmWalletDepositFromBankTransaction(id)\n"
        "  FirebaseDataService -> PayosService: getPaymentLink(paymentLinkId)\n"
        "  PayosService -> PayOS API: GET /v2/payment-requests/{id}\n"
        "  PayOS API --> PayosService: status\n"
        "end\n"
        "alt status == PAID\n"
        "  FirebaseDataService -> Firestore.transaction: complete walletRequest + increment walletBalance\n"
        "  ProfileScreen -> User: Hiển thị hiệu ứng thanh toán thành công\n"
        "else Chưa thanh toán\n"
        "  ProfileScreen -> User: Thông báo tiếp tục kiểm tra\n"
        "end",
    )
    image_note(doc, "Chèn ảnh màn hình QR PayOS, ảnh giao dịch thành công và ảnh hiệu ứng thanh toán thành công.")

    heading(doc, "3.11.6 Sequence Diagram: Rút tiền từ ví", 2)
    code_block(
        doc,
        "User -> ProfileScreen: Chọn Rút tiền\n"
        "ProfileScreen -> FirebaseDataService: requestWalletWithdrawal(amount)\n"
        "FirebaseDataService -> Firestore.users: Lấy walletBalance và thông tin ngân hàng\n"
        "alt Hợp lệ\n"
        "  FirebaseDataService -> Firestore.walletRequests: Tạo withdrawal pending\n"
        "  FirebaseDataService -> Firestore.users: Tạm trừ walletBalance\n"
        "  FirebaseDataService -> Admin: Tạo thông báo yêu cầu rút\n"
        "else Không hợp lệ\n"
        "  ProfileScreen -> User: Thông báo lỗi\n"
        "end\n"
        "Admin -> AdminDashboardScreen: Xử lý yêu cầu\n"
        "AdminDashboardScreen -> FirebaseDataService: completeWalletWithdrawal(requestId)\n"
        "FirebaseDataService -> Firestore.walletRequests: status = completed\n"
        "FirebaseDataService -> User: Tạo thông báo rút tiền hoàn tất",
    )
    image_note(doc, "Chèn ảnh màn hình form rút tiền và ảnh admin xử lý yêu cầu rút tiền.")

    heading(doc, "3.11.7 Sequence Diagram: Chat hỗ trợ admin", 2)
    code_block(
        doc,
        "User -> ProfileScreen: Bấm Chat admin\n"
        "ProfileScreen -> FirebaseDataService: ensureAdminConversation(topic)\n"
        "FirebaseDataService -> Firestore.conversations: Tạo/lấy conversation\n"
        "ProfileScreen -> ChatScreen: Mở cuộc trò chuyện\n"
        "User -> ChatScreen: Nhập tin nhắn\n"
        "ChatScreen -> FirebaseDataService: sendChatMessage(conversationId, text)\n"
        "FirebaseDataService -> Firestore.messages: Thêm message\n"
        "FirebaseDataService -> Firestore.conversations: Cập nhật lastMessage\n"
        "FirebaseDataService -> notifications: Tạo thông báo cho người nhận\n"
        "alt Chat hỗ trợ admin\n"
        "  FirebaseDataService -> SupportBotService: replyFor(text)\n"
        "  SupportBotService --> FirebaseDataService: Nội dung phản hồi\n"
        "  FirebaseDataService -> Firestore.messages: Thêm tin nhắn bot\n"
        "end",
    )
    image_note(doc, "Chèn ảnh màn hình danh sách chat và màn hình chat với admin.")


def chapter_4(doc):
    heading(doc, "Chương 4. THIẾT KẾ DỮ LIỆU", 1)
    heading(doc, "4.1 Mô hình dữ liệu tổng quan", 2)
    paragraph(doc, "Dữ liệu của EduShare được lưu trên Cloud Firestore theo dạng collection/document. Mỗi collection đại diện cho một nhóm dữ liệu nghiệp vụ.")
    add_table(
        doc,
        ["Collection", "Mô tả"],
        [
            ["users", "Thông tin tài khoản, hồ sơ, ví, địa chỉ, quyền admin"],
            ["products", "Thông tin sản phẩm được đăng bán"],
            ["favorites", "Danh sách sản phẩm yêu thích của người dùng"],
            ["orders", "Thông tin đơn hàng và trạng thái xử lý"],
            ["walletRequests", "Yêu cầu nạp tiền/rút tiền của người dùng"],
            ["notifications", "Thông báo gửi đến người dùng"],
            ["conversations", "Cuộc trò chuyện giữa người dùng và admin/người bán"],
            ["bankTransactions", "Thông tin giao dịch ngân hàng dùng cho đối soát khi có backend hỗ trợ"],
        ],
    )

    heading(doc, "4.2 Mô tả chi tiết một số collection", 2)
    add_table(
        doc,
        ["Collection", "Trường dữ liệu tiêu biểu"],
        [
            ["users", "name, email, phone, university, walletBalance, isAdmin, shippingAddress, bankName"],
            ["products", "title, author, category, price, stockQuantity, sellerUid, imageUrl, university"],
            ["orders", "buyerUid, sellerUid, productId, quantity, totalPrice, status, paymentMethod, createdAt"],
            ["walletRequests", "userUid, type, requestedAmount, creditedAmount, status, payosOrderCode, payosStatus"],
            ["notifications", "userUid, title, body, type, isRead, createdAt"],
            ["conversations", "participantIds, lastMessage, updatedAt; subcollection messages"],
        ],
    )

    heading(doc, "4.2.1 Collection users", 2)
    add_table(
        doc,
        ["Trường", "Kiểu dữ liệu", "Mô tả"],
        [
            ["name", "String", "Họ tên người dùng"],
            ["email", "String", "Email đăng nhập"],
            ["phone", "String", "Số điện thoại liên hệ"],
            ["university", "String", "Trường học"],
            ["walletBalance", "Number", "Số dư ví EduShare"],
            ["isAdmin", "Boolean", "Đánh dấu tài khoản admin"],
            ["shippingAddress", "String", "Địa chỉ nhận hàng"],
            ["shippingLatitude/shippingLongitude", "Number", "Tọa độ giao hàng"],
            ["bankName/bankAccountNumber", "String", "Thông tin ngân hàng để rút tiền"],
        ],
    )

    heading(doc, "4.2.2 Collection products", 2)
    add_table(
        doc,
        ["Trường", "Kiểu dữ liệu", "Mô tả"],
        [
            ["title", "String", "Tên sản phẩm"],
            ["author", "String", "Tác giả hoặc mô tả ngắn"],
            ["category", "String", "Danh mục sản phẩm"],
            ["price", "Number", "Giá bán"],
            ["stockQuantity", "Number", "Số lượng tồn kho"],
            ["sellerUid", "String", "ID người bán"],
            ["imageUrl", "String", "Đường dẫn hoặc dữ liệu ảnh sản phẩm"],
            ["createdAt", "Date/String", "Thời điểm đăng sản phẩm"],
        ],
    )

    heading(doc, "4.2.3 Collection orders", 2)
    add_table(
        doc,
        ["Trường", "Kiểu dữ liệu", "Mô tả"],
        [
            ["buyerUid", "String", "ID người mua"],
            ["sellerUid", "String", "ID người bán"],
            ["productId", "String", "ID sản phẩm"],
            ["quantity", "Number", "Số lượng mua"],
            ["totalPrice", "Number", "Tổng tiền của dòng đơn hàng"],
            ["status", "String", "Trạng thái đơn hàng"],
            ["paymentMethod", "String", "Phương thức thanh toán"],
            ["sellerPayoutAmount", "Number", "Số tiền giải ngân cho người bán"],
            ["platformFeeAmount", "Number", "Phí nền tảng"],
        ],
    )

    heading(doc, "4.2.4 Collection walletRequests", 2)
    add_table(
        doc,
        ["Trường", "Kiểu dữ liệu", "Mô tả"],
        [
            ["userUid", "String", "ID người tạo yêu cầu"],
            ["type", "String", "deposit hoặc withdrawal"],
            ["requestedAmount", "Number", "Số tiền người dùng nạp/rút"],
            ["creditedAmount", "Number", "Số tiền thực cộng vào ví"],
            ["status", "String", "pending, completed hoặc cancelled"],
            ["transferNote", "String", "Nội dung giao dịch"],
            ["payosOrderCode", "Number", "Mã đơn PayOS"],
            ["payosPaymentLinkId", "String", "ID link thanh toán PayOS"],
            ["payosStatus", "String", "Trạng thái PayOS"],
        ],
    )

    heading(doc, "4.2.5 Collection conversations và messages", 2)
    paragraph(doc, "Mỗi document trong conversations đại diện cho một cuộc trò chuyện. Tin nhắn được lưu trong subcollection messages để dễ truy vấn theo từng cuộc hội thoại.")
    bullet(doc, "conversations.participantIds: danh sách ID người tham gia.")
    bullet(doc, "conversations.lastMessage: tin nhắn cuối cùng để hiển thị ở danh sách chat.")
    bullet(doc, "messages.senderUid: người gửi tin nhắn.")
    bullet(doc, "messages.text: nội dung tin nhắn.")
    bullet(doc, "messages.createdAt: thời điểm gửi.")

    heading(doc, "4.3 Tối ưu và an toàn dữ liệu", 2)
    bullet(doc, "Sử dụng Firebase Authentication để xác định request.auth.uid khi đọc/ghi dữ liệu.")
    bullet(doc, "Các thao tác cộng/trừ ví quan trọng sử dụng transaction hoặc FieldValue.increment nhằm hạn chế sai lệch số dư.")
    bullet(doc, "Dữ liệu sản phẩm và đơn hàng được tách thành model riêng giúp dễ kiểm soát kiểu dữ liệu.")
    bullet(doc, "Các collection cần được bảo vệ bằng Firestore Security Rules để giới hạn quyền user và admin.")

    heading(doc, "4.4 Quan hệ dữ liệu logic", 2)
    paragraph(doc, "Mặc dù Firestore là NoSQL và không bắt buộc khóa ngoại như cơ sở dữ liệu quan hệ, hệ thống vẫn duy trì các quan hệ logic thông qua ID:")
    bullet(doc, "orders.buyerUid liên kết đến users.id của người mua.")
    bullet(doc, "orders.sellerUid liên kết đến users.id của người bán.")
    bullet(doc, "orders.productId liên kết đến products.id.")
    bullet(doc, "walletRequests.userUid liên kết đến users.id.")
    bullet(doc, "favorites.userUid và favorites.productId liên kết người dùng với sản phẩm yêu thích.")
    bullet(doc, "notifications.userUid xác định người nhận thông báo.")

    heading(doc, "4.5 Các ràng buộc nghiệp vụ", 2)
    bullet(doc, "Người dùng phải đăng nhập mới được đặt hàng, yêu thích, chat hoặc sử dụng ví.")
    bullet(doc, "Không cho thanh toán bằng ví nếu walletBalance nhỏ hơn tổng tiền đơn hàng.")
    bullet(doc, "Không cho đặt hàng nếu số lượng sản phẩm vượt quá tồn kho hiện tại.")
    bullet(doc, "Yêu cầu rút tiền chỉ được tạo khi người dùng có tài khoản ngân hàng và số dư đủ.")
    bullet(doc, "Yêu cầu nạp ví chỉ được cộng tiền khi trạng thái thanh toán hợp lệ.")


def chapter_5(doc):
    heading(doc, "Chương 5. THIẾT KẾ GIAO DIỆN", 1)
    heading(doc, "5.1 Nguyên tắc thiết kế UI/UX", 2)
    paragraph(doc, "Giao diện EduShare hướng đến sự đơn giản, rõ ràng và phù hợp với người dùng sinh viên. Các chức năng quan trọng như tìm kiếm, giỏ hàng, ví và hồ sơ được đặt ở thanh điều hướng dưới cùng để dễ thao tác.")
    bullet(doc, "Màu chủ đạo xanh ngọc tạo cảm giác hiện đại và thân thiện.")
    bullet(doc, "Các card sản phẩm, card ví và card thống kê giúp thông tin dễ quét.")
    bullet(doc, "Hiệu ứng chuyển tab, hiệu ứng thanh toán thành công và thông báo trạng thái giúp tăng phản hồi trực quan.")

    heading(doc, "5.2 Các giao diện chính", 2)
    add_table(
        doc,
        ["Giao diện", "File code", "Mô tả"],
        [
            ["Trang điều hướng chính", "lib/screens/app_shell.dart", "Quản lý tab Trang chủ, Tìm kiếm, Giỏ hàng, Hồ sơ"],
            ["Trang chủ", "lib/screens/home_screen.dart", "Hiển thị sản phẩm nổi bật, sản phẩm mới và gợi ý"],
            ["Tìm kiếm", "lib/screens/search_screen.dart", "Tìm sản phẩm theo từ khóa"],
            ["Giỏ hàng", "lib/screens/cart_screen.dart", "Quản lý sản phẩm trong giỏ và thanh toán"],
            ["Hồ sơ/Ví", "lib/screens/profile_screen.dart", "Quản lý hồ sơ, ví EduShare, nạp/rút tiền"],
            ["Admin", "lib/screens/admin_dashboard_screen.dart", "Quản trị người dùng, sản phẩm, đơn hàng và ví"],
            ["Chat", "lib/screens/chat_screen.dart", "Trao đổi tin nhắn hỗ trợ"],
        ],
    )

    heading(doc, "5.2.1 Màn hình hồ sơ và ví EduShare", 2)
    paragraph(doc, "ProfileScreen là một trong những màn hình quan trọng nhất vì tập trung thông tin cá nhân, ví EduShare, tài khoản rút tiền, lịch sử hoạt động và đường dẫn đến admin dashboard nếu người dùng là admin. Card ví hiển thị số dư hiện tại, nút nạp tiền và nút rút tiền.")
    paragraph(doc, "Chức năng nạp tiền được mở bằng _showWalletDepositSheet. Sau khi người dùng nhập số tiền, hệ thống tạo WalletRequest và chuyển sang _WalletTopupScreen để hiển thị QR/thông tin thanh toán.")
    image_note(doc, "Chèn ảnh màn hình Hồ sơ có card ví EduShare, số dư ví và nút Nạp tiền/Rút tiền.")
    image_note(doc, "Chèn ảnh màn hình QR nạp tiền PayOS hoặc màn hình xác nhận nạp ví thành công.")

    heading(doc, "5.2.2 Màn hình giỏ hàng", 2)
    paragraph(doc, "CartScreen hiển thị danh sách sản phẩm đã chọn, tổng số lượng, tổng tiền và các phương thức thanh toán. Sau khi gỡ bỏ chuyển khoản QR tự động, màn hình tập trung vào hai phương thức: thanh toán khi nhận hàng và thanh toán bằng ví EduShare.")
    image_note(doc, "Chèn ảnh giỏ hàng có ít nhất 2 sản phẩm, hiển thị tổng tiền và số lượng.")
    image_note(doc, "Chèn ảnh bottom sheet chọn thanh toán gồm COD và ví EduShare.")

    heading(doc, "5.2.3 Màn hình admin", 2)
    paragraph(doc, "AdminDashboardScreen cung cấp giao diện quản trị gồm danh sách người dùng, sản phẩm, đơn hàng, yêu cầu nạp/rút ví. Admin có thể xác nhận thanh toán, xử lý rút tiền, giải ngân cho người bán và quản lý sản phẩm vi phạm.")
    image_note(doc, "Chèn ảnh dashboard admin phần yêu cầu nạp/rút ví hoặc danh sách đơn hàng.")

    heading(doc, "5.2.4 Màn hình chat", 2)
    paragraph(doc, "ChatScreen hiển thị tin nhắn theo thời gian thực từ Firestore. Người dùng có thể gửi tin nhắn cho admin để nhận hỗ trợ về tài khoản, ví, đơn hàng hoặc sản phẩm.")
    image_note(doc, "Chèn ảnh màn hình Chat admin với ít nhất một tin nhắn người dùng và một tin nhắn phản hồi.")

    heading(doc, "5.2.5 Các vị trí cần bổ sung hình ảnh sản phẩm", 2)
    paragraph(doc, "Để báo cáo trực quan hơn, cần bổ sung các ảnh sản phẩm thật hoặc ảnh chụp màn hình có sản phẩm tại các vị trí sau:")
    add_table(
        doc,
        ["Vị trí trong báo cáo", "Loại ảnh cần thêm", "Mục đích"],
        [
            ["Chương 1 - Mô tả sản phẩm", "Ảnh tổng quan app EduShare có danh sách sản phẩm", "Giúp người đọc thấy ngay ứng dụng đang giải quyết bài toán mua bán sản phẩm học tập"],
            ["Chương 3 - Đăng bán sản phẩm", "Ảnh form đăng sản phẩm có ảnh bìa sách/dụng cụ", "Minh họa dữ liệu đầu vào khi người bán đăng sản phẩm"],
            ["Chương 3 - Thêm vào giỏ hàng", "Ảnh card sản phẩm và giỏ hàng sau khi thêm", "Minh họa luồng mua hàng"],
            ["Chương 4 - Collection products", "Ảnh ví dụ sản phẩm thật: sách, máy tính, dụng cụ học tập", "Liên hệ trường dữ liệu sản phẩm với giao diện thực tế"],
            ["Chương 5 - Trang chủ", "Ảnh danh sách sản phẩm mới/nổi bật", "Minh họa giao diện chính"],
            ["Chương 5 - Tìm kiếm", "Ảnh kết quả tìm kiếm theo tên sách hoặc danh mục", "Minh họa chức năng tìm kiếm"],
            ["Chương 5 - Chi tiết sản phẩm", "Ảnh màn hình chi tiết sản phẩm", "Thể hiện thông tin giá, tồn kho, người bán, ảnh sản phẩm"],
            ["Phụ lục", "Ảnh cấu trúc thư mục assets/images hoặc ví dụ ảnh sản phẩm", "Chứng minh dữ liệu ảnh được dùng trong ứng dụng"],
        ],
    )
    image_note(doc, "Chèn ít nhất 5 ảnh sản phẩm thật: sách giáo trình, sách tham khảo, máy tính, dụng cụ học tập, combo tài liệu.")

    heading(doc, "5.3 Hiệu ứng giao diện", 2)
    paragraph(doc, "Ứng dụng có một số hiệu ứng nhằm tăng trải nghiệm người dùng:")
    bullet(doc, "Hiệu ứng chuyển tab sử dụng AnimatedOpacity, AnimatedSlide và AnimatedScale trong AppShell.")
    bullet(doc, "Hiệu ứng thanh toán thành công trong ProfileScreen gồm overlay, vòng sáng, biểu tượng xác nhận và CustomPainter tạo hiệu ứng tia sáng.")
    bullet(doc, "Các nút điều hướng và card có AnimatedContainer/AnimatedScale giúp thao tác mượt hơn.")

    heading(doc, "5.4 Trải nghiệm người dùng khi thanh toán", 2)
    paragraph(doc, "Khi thanh toán hoặc nạp ví, người dùng cần được phản hồi rõ ràng để tránh nhầm lẫn. EduShare sử dụng loading indicator khi đang kiểm tra giao dịch, snackbar khi chưa tìm thấy giao dịch và overlay thành công khi tiền đã được cộng vào ví.")
    paragraph(doc, "Hiệu ứng thành công trong ProfileScreen sử dụng CustomPainter để vẽ các tia sáng xung quanh biểu tượng xác nhận. Đây là cách tạo hiệu ứng trực tiếp bằng Flutter mà không cần thêm thư viện animation bên ngoài.")

    heading(doc, "5.5 Màu sắc và phong cách", 2)
    add_table(
        doc,
        ["Thành phần", "Màu/Phong cách"],
        [
            ["Màu chính", "Xanh ngọc AppColors.primary"],
            ["Nền", "Xám rất nhạt AppColors.bg"],
            ["Trạng thái cảnh báo", "Amber"],
            ["Trạng thái lỗi", "Đỏ"],
            ["Card", "Nền trắng, bo góc, shadow nhẹ"],
            ["Thanh điều hướng", "BottomAppBar có notch cho nút đăng sản phẩm"],
        ],
    )

    heading(doc, "5.6 Khả năng responsive", 2)
    paragraph(doc, "Ứng dụng được xây dựng bằng Flutter nên có thể chạy trên nhiều kích thước màn hình. Các màn hình sử dụng ListView, Expanded, Flexible, Wrap và các constraint phù hợp để tránh tràn nội dung. Các form như nạp tiền, rút tiền, chọn thanh toán được đặt trong bottom sheet để phù hợp với thao tác trên điện thoại.")


def chapter_6(doc):
    heading(doc, "Chương 6. KẾT LUẬN", 1)
    heading(doc, "6.1 Những kết quả đạt được", 2)
    for t in [
        "Xây dựng được ứng dụng Flutter có đầy đủ luồng đăng nhập, hồ sơ, sản phẩm, giỏ hàng và đơn hàng.",
        "Tích hợp Cloud Firestore để lưu trữ dữ liệu sản phẩm, người dùng, ví, đơn hàng, chat và thông báo.",
        "Xây dựng chức năng ví EduShare với nạp tiền, rút tiền và thanh toán bằng ví.",
        "Tích hợp PayOS để hỗ trợ tạo thanh toán và kiểm tra trạng thái nạp tiền.",
        "Xây dựng trang admin để quản lý người dùng, sản phẩm, yêu cầu ví và đơn hàng.",
        "Cải thiện trải nghiệm bằng hiệu ứng chuyển tab và hiệu ứng thanh toán thành công.",
    ]:
        bullet(doc, t)

    heading(doc, "6.2 Hạn chế còn tồn tại", 2)
    for t in [
        "Một số chức năng thanh toán vẫn cần backend riêng để bảo mật tuyệt đối khóa API và xử lý webhook PayOS.",
        "Chưa có hệ thống đánh giá người bán/người mua.",
        "Chưa có bộ lọc nâng cao theo khoảng giá, vị trí hoặc độ uy tín.",
        "Chưa có thống kê doanh thu chuyên sâu cho admin và người bán.",
    ]:
        bullet(doc, t)

    heading(doc, "6.2.1 Hạn chế về bảo mật thanh toán", 2)
    paragraph(doc, "Trong phiên bản demo, PayOS được gọi từ ứng dụng Flutter thông qua dart-define. Cách này thuận tiện cho kiểm thử nhưng chưa phải giải pháp tối ưu cho môi trường production, vì khóa API phía client có thể bị trích xuất. Hướng triển khai tốt hơn là đưa logic tạo link, xác minh webhook và cộng tiền ví vào Firebase Cloud Functions hoặc backend riêng.")

    heading(doc, "6.2.2 Hạn chế về kiểm thử", 2)
    paragraph(doc, "Đề tài hiện chủ yếu kiểm thử thủ công qua emulator và thao tác trực tiếp trên ứng dụng. Trong tương lai cần bổ sung unit test cho service, widget test cho giao diện và integration test cho các luồng chính như đăng nhập, đặt hàng, nạp ví.")

    heading(doc, "6.3 Hướng phát triển trong tương lai", 2)
    for t in [
        "Triển khai Firebase Cloud Functions để xử lý webhook PayOS và bảo vệ API key.",
        "Bổ sung chức năng đánh giá, bình luận và báo cáo sản phẩm vi phạm.",
        "Tối ưu gợi ý sản phẩm dựa trên hành vi người dùng.",
        "Bổ sung dashboard thống kê doanh thu, đơn hàng và lượt bán.",
        "Hoàn thiện kiểm thử tự động và quy trình triển khai.",
    ]:
        bullet(doc, t)

    heading(doc, "6.4 Tổng kết", 2)
    paragraph(
        doc,
        "EduShare đã hoàn thành mục tiêu xây dựng một ứng dụng hỗ trợ sinh viên mua bán và trao đổi sách, dụng cụ học tập với các chức năng tương đối đầy đủ. Thông qua đề tài, em đã vận dụng được kiến thức về Flutter, Firebase, mô hình dữ liệu NoSQL, quản lý trạng thái, tích hợp thanh toán và thiết kế giao diện người dùng. Sản phẩm có thể tiếp tục phát triển thành nền tảng thương mại học đường hoàn chỉnh nếu được bổ sung backend bảo mật, kiểm thử tự động và các chức năng đánh giá/ngăn chặn gian lận.",
    )


def chapter_8_appendix(doc):
    heading(doc, "PHỤ LỤC. MÔ TẢ MỘT SỐ ĐOẠN CODE QUAN TRỌNG", 1)
    heading(doc, "A.1 FirebaseDataService", 2)
    paragraph(doc, "FirebaseDataService là lớp trung tâm của hệ thống. Lớp này định nghĩa các collection Firestore và cung cấp các hàm thao tác dữ liệu cho toàn bộ ứng dụng.")
    bullet(doc, "ensureUserProfile: đảm bảo mỗi tài khoản Firebase Auth có một hồ sơ người dùng tương ứng trong collection users.")
    bullet(doc, "getAllProducts/getRecentProducts/searchProducts: phục vụ hiển thị và tìm kiếm sản phẩm.")
    bullet(doc, "createOrdersFromCart: tạo đơn hàng từ giỏ hàng và cập nhật tồn kho.")
    bullet(doc, "createWalletPaidOrdersFromCart: tạo đơn thanh toán bằng ví và trừ số dư.")
    bullet(doc, "requestWalletDeposit/requestWalletWithdrawal: tạo yêu cầu nạp/rút ví.")
    bullet(doc, "_completeWalletDeposit: hoàn tất nạp ví bằng transaction và FieldValue.increment.")

    heading(doc, "A.2 PayosService", 2)
    paragraph(doc, "PayosService chịu trách nhiệm giao tiếp với PayOS. Lớp này tạo chữ ký HMAC SHA256, gửi request tạo payment link và gọi API lấy trạng thái thanh toán.")
    bullet(doc, "createWalletTopupLink: tạo link thanh toán cho yêu cầu nạp ví.")
    bullet(doc, "getPaymentLink: kiểm tra trạng thái thanh toán dựa trên paymentLinkId hoặc orderCode.")
    bullet(doc, "_signature: tạo chữ ký từ các tham số amount, cancelUrl, description, orderCode, returnUrl.")

    heading(doc, "A.3 CartProvider", 2)
    paragraph(doc, "CartProvider quản lý danh sách sản phẩm trong giỏ hàng. Đây là lớp state management giúp CartScreen tự động cập nhật khi người dùng thêm, xóa hoặc thay đổi số lượng sản phẩm.")

    heading(doc, "A.4 AppShell", 2)
    paragraph(doc, "AppShell là màn hình khung sau khi đăng nhập. Nó quản lý bottom navigation, nút đăng sản phẩm và hiệu ứng chuyển tab. Hàm _buildAnimatedTabBody dùng Stack kết hợp AnimatedOpacity, AnimatedSlide và AnimatedScale để tạo chuyển cảnh mượt.")

    heading(doc, "A.5 ProfileScreen", 2)
    paragraph(doc, "ProfileScreen chứa nhiều chức năng liên quan đến hồ sơ và ví. Các hàm quan trọng gồm _walletCard, _showWalletDepositSheet, _showWalletWithdrawSheet, _checkBankPayment và _showPaymentSuccessEffect.")

    heading(doc, "A.6 AdminDashboardScreen", 2)
    paragraph(doc, "AdminDashboardScreen tập hợp các thao tác quản trị. Giao diện này gọi các hàm trong FirebaseDataService để lấy danh sách người dùng, sản phẩm, đơn hàng, yêu cầu ví và thực hiện các thao tác admin.")


def testing_chapter(doc):
    heading(doc, "Chương 7. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    heading(doc, "7.1 Mục tiêu kiểm thử", 2)
    paragraph(doc, "Kiểm thử nhằm đảm bảo các chức năng chính của EduShare hoạt động đúng, dữ liệu được ghi nhận đầy đủ và giao diện phản hồi phù hợp với thao tác người dùng.")
    heading(doc, "7.2 Các trường hợp kiểm thử tiêu biểu", 2)
    add_table(
        doc,
        ["Mã TC", "Chức năng", "Dữ liệu/Thao tác", "Kết quả mong đợi"],
        [
            ["TC01", "Đăng nhập", "Nhập email/mật khẩu hợp lệ", "Đăng nhập thành công và vào AppShell"],
            ["TC02", "Đăng bán sản phẩm", "Nhập đầy đủ tên, giá, số lượng", "Sản phẩm được lưu vào products"],
            ["TC03", "Tìm kiếm", "Nhập từ khóa tên sách", "Danh sách sản phẩm phù hợp được hiển thị"],
            ["TC04", "Thêm giỏ hàng", "Bấm thêm sản phẩm", "Số lượng giỏ hàng tăng"],
            ["TC05", "Thanh toán bằng ví đủ tiền", "Chọn ví EduShare", "Đơn hàng được tạo, ví bị trừ"],
            ["TC06", "Thanh toán bằng ví thiếu tiền", "Số dư nhỏ hơn tổng tiền", "Hiển thị thông báo số dư không đủ"],
            ["TC07", "Nạp ví", "Tạo yêu cầu nạp tiền", "walletRequests được tạo với trạng thái pending"],
            ["TC08", "PayOS PAID", "Giao dịch trả về PAID", "Ví được cộng tiền và yêu cầu chuyển completed"],
            ["TC09", "Chat admin", "Gửi tin nhắn hỗ trợ", "Tin nhắn được lưu và hiển thị realtime"],
            ["TC10", "Admin duyệt rút tiền", "Admin xử lý yêu cầu rút", "Yêu cầu chuyển completed và có thông báo"],
        ],
    )

    heading(doc, "7.3 Đánh giá kết quả", 2)
    paragraph(doc, "Qua quá trình kiểm thử thủ công trên emulator, các luồng chính như đăng nhập, xem sản phẩm, giỏ hàng, ví và chat hoạt động đúng theo mục tiêu đề ra. Một số chức năng liên quan thanh toán cần tiếp tục hoàn thiện bằng backend để đảm bảo bảo mật và tự động hóa hoàn toàn trong môi trường thực tế.")


def references(doc):
    heading(doc, "Chương 8. TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Flutter Documentation: https://docs.flutter.dev/",
        "Dart Documentation: https://dart.dev/guides",
        "Firebase Documentation: https://firebase.google.com/docs",
        "Cloud Firestore Documentation: https://firebase.google.com/docs/firestore",
        "PayOS Documentation: https://payos.vn/docs/",
        "Provider package documentation: https://pub.dev/packages/provider",
        "Flutter Map package documentation: https://pub.dev/packages/flutter_map",
    ]
    for ref in refs:
        paragraph(doc, ref)


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    acknowledgements(doc)
    toc(doc)
    chapter_1(doc)
    page_break(doc)
    chapter_2(doc)
    page_break(doc)
    chapter_3(doc)
    page_break(doc)
    chapter_4(doc)
    page_break(doc)
    chapter_5(doc)
    page_break(doc)
    chapter_6(doc)
    page_break(doc)
    testing_chapter(doc)
    page_break(doc)
    chapter_8_appendix(doc)
    page_break(doc)
    references(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

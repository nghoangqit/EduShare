from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Bao_cao_DACS3_EduShare_co_muc_luc.docx"


def set_run(run, size=13, bold=False, italic=False, color=None, font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def p(doc, text="", style=None, align=None, size=13, bold=False, italic=False):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return para


def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    if para.runs:
        set_run(para.runs[0], size=15 if level == 1 else 14, bold=True)
    return para


def bullet(doc, text):
    return p(doc, text, style="List Bullet")


def num(doc, text):
    return p(doc, text, style="List Number")


def code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.7)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run(run, size=10.5, font="Courier New")
    return para


def image_note(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(f"[[CẦN THÊM HÌNH ẢNH: {text}]]")
    set_run(run, size=12, bold=True, italic=True, color=(192, 0, 0))
    return para


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run(run, size=12, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[i].paragraphs:
                for run in para.runs:
                    set_run(run, size=12)
    doc.add_paragraph()
    return t


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    for name in ["Normal", "List Bullet", "List Number"]:
        doc.styles[name].font.name = "Times New Roman"
        doc.styles[name].font.size = Pt(13)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.name = "Times New Roman"
        doc.styles[name].font.color.rgb = RGBColor(0, 0, 0)


def page(doc):
    doc.add_page_break()


def cover(doc):
    p(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN &", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    p(doc, "TRUYỀN THÔNG VIỆT HÀN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
    p(doc, "Khoa Khoa Học Máy Tính", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    p(doc, "ĐỒ ÁN CƠ SỞ 3", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    doc.add_paragraph()
    p(doc, "XÂY DỰNG ỨNG DỤNG MOBILE EDUSHARE HỖ TRỢ MUA BÁN SÁCH VÀ DỤNG CỤ HỌC TẬP", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=17)
    doc.add_paragraph()
    p(doc, "Sinh viên thực hiện:      Trần Ngọc Hoàng", size=13)
    p(doc, "Lớp:                      ........................................", size=13)
    p(doc, "MSSV:                     24ITB056", size=13)
    p(doc, "Giảng viên hướng dẫn:     ........................................", size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    p(doc, "Đà Nẵng, tháng 5 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
    page(doc)


def preface(doc):
    h(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN", 1)
    for _ in range(16):
        p(doc, "........................................................................................................................")
    page(doc)
    h(doc, "LỜI CẢM ƠN", 1)
    texts = [
        "Trước tiên, em xin gửi lời cảm ơn chân thành đến quý thầy cô đã hướng dẫn, góp ý và tạo điều kiện để em hoàn thành đồ án cơ sở 3.",
        "Em xin cảm ơn Khoa Khoa Học Máy Tính đã cung cấp nền tảng kiến thức về lập trình ứng dụng di động, cơ sở dữ liệu, phân tích thiết kế hệ thống và triển khai phần mềm.",
        "Trong quá trình thực hiện đề tài EduShare, em đã có cơ hội vận dụng kiến thức Flutter, Firebase, quản lý trạng thái, thiết kế giao diện, tích hợp thanh toán PayOS và xử lý dữ liệu thời gian thực.",
        "Do thời gian và kinh nghiệm còn hạn chế, báo cáo cũng như sản phẩm chắc chắn vẫn còn thiếu sót. Em rất mong nhận được ý kiến đóng góp từ quý thầy cô để hệ thống được hoàn thiện hơn.",
        "Em xin trân trọng cảm ơn!",
    ]
    for text in texts:
        p(doc, text)
    doc.add_paragraph()
    p(doc, "Sinh viên,", align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(doc, "Trần Ngọc Hoàng", align=WD_ALIGN_PARAGRAPH.RIGHT)
    page(doc)


def lists(doc):
    h(doc, "MỤC LỤC", 1)
    toc_items = [
        "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN",
        "LỜI CẢM ƠN",
        "DANH MỤC CÁC TỪ VIẾT TẮT",
        "DANH MỤC HÌNH ẢNH",
        "CHƯƠNG 1 GIỚI THIỆU ĐỀ TÀI",
        "1.1. Lý do chọn đề tài",
        "1.1.1. Nhu cầu trao đổi tài liệu học tập trong sinh viên",
        "1.1.2. Hạn chế của các kênh mua bán truyền thống",
        "1.1.3. Giải pháp đề xuất",
        "1.2. Mục tiêu của đề tài",
        "1.2.1. Mục tiêu tổng quát",
        "1.2.2. Mục tiêu cụ thể",
        "1.3. Ý nghĩa thực tiễn",
        "1.4. Mô tả hệ thống",
        "1.5. Công nghệ sử dụng",
        "CHƯƠNG 2 CƠ SỞ LÝ THUYẾT",
        "2.1. Framework Flutter và ngôn ngữ lập trình Dart",
        "2.2. Firebase Backend-as-a-Service",
        "2.3. Cloud Firestore và mô hình NoSQL",
        "2.4. Quản lý trạng thái với Provider",
        "2.5. Thanh toán điện tử và PayOS",
        "2.6. Bảo mật dữ liệu trong ứng dụng mobile",
        "2.7. UML và mô hình hóa hệ thống",
        "2.8. Mô hình C4",
        "CHƯƠNG 3 PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "3.1. Phân tích yêu cầu hệ thống",
        "3.2. Sơ đồ Use-case",
        "3.3. Kiến trúc hệ thống theo mô hình C4",
        "3.4. Thiết kế chi tiết cấu trúc hệ thống lớp",
        "3.5. Thiết kế các sơ đồ hành vi",
        "3.6. Sơ đồ giao tiếp",
        "3.7. Sơ đồ tuần tự chi tiết",
        "3.8. Thiết kế cấu trúc cơ sở dữ liệu chi tiết",
        "CHƯƠNG 4 TRIỂN KHAI VÀ XÂY DỰNG ỨNG DỤNG",
        "4.1. Môi trường phát triển và tổ chức mã nguồn",
        "4.2. Cấu trúc thư mục dự án",
        "4.3. Triển khai lớp nghiệp vụ",
        "4.4. Triển khai lớp hiển thị và quản lý trạng thái",
        "4.5. Giao diện ứng dụng thực tế",
        "CHƯƠNG 5 KIỂM THỬ VÀ ĐÁNH GIÁ HỆ THỐNG",
        "5.1. Kiểm thử hệ thống",
        "5.2. Kết quả đạt được",
        "5.3. Đánh giá hệ thống",
        "CHƯƠNG 6 KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        "6.1. Tóm tắt nội dung chính của đề tài",
        "6.2. Định hướng mở rộng trong tương lai",
        "TÀI LIỆU THAM KHẢO",
    ]
    for item in toc_items:
        p(doc, item)
    page(doc)

    h(doc, "DANH MỤC CÁC TỪ VIẾT TẮT", 1)
    table(doc, ["Từ viết tắt", "Ý nghĩa"], [
        ["API", "Application Programming Interface"],
        ["BaaS", "Backend as a Service"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["QR", "Quick Response Code"],
        ["UI/UX", "User Interface / User Experience"],
        ["UML", "Unified Modeling Language"],
        ["NoSQL", "Cơ sở dữ liệu phi quan hệ"],
        ["SDK", "Software Development Kit"],
    ])
    h(doc, "DANH MỤC HÌNH ẢNH", 1)
    for item in [
        "Hình 3.2 – Sơ đồ Use-case tổng quan",
        "Hình 3.3.1 – Sơ đồ ngữ cảnh (Context Diagram)",
        "Hình 3.3.2 – Sơ đồ Container (Container Diagram)",
        "Hình 3.3.3 – Sơ đồ thành phần (Component Diagram)",
        "Hình 3.4 – Sơ đồ lớp (Class Diagram)",
        "Hình 3.5.1 – Sơ đồ hoạt động (Activity Diagram)",
        "Hình 3.5.2 – Sơ đồ trạng thái (State Machine Diagram)",
        "Hình 3.6 – Sơ đồ giao tiếp (Communication Diagram)",
        "Hình 3.7 – Sơ đồ tuần tự (Sequence Diagram)",
        "Hình 3.8 – Sơ đồ mối quan hệ thực thể ERD",
        "Hình 4.5.1 – Màn hình Trang chủ",
        "Hình 4.5.2 – Màn hình Chi tiết sản phẩm",
        "Hình 4.5.3 – Màn hình Giỏ hàng và thanh toán",
        "Hình 4.5.4 – Màn hình Ví EduShare và PayOS",
        "Hình 4.5.5 – Màn hình Admin Dashboard",
    ]:
        p(doc, item)
    page(doc)


def chapter1(doc):
    h(doc, "CHƯƠNG 1 GIỚI THIỆU ĐỀ TÀI", 1)
    h(doc, "1.1. Lý do chọn đề tài", 1)
    h(doc, "1.1.1. Nhu cầu trao đổi tài liệu học tập trong sinh viên", 2)
    p(doc, "Trong môi trường đại học, sinh viên thường xuyên cần mua sách, giáo trình, tài liệu tham khảo, máy tính và các dụng cụ học tập. Tuy nhiên, không phải sinh viên nào cũng có đủ điều kiện mua toàn bộ sản phẩm mới. Bên cạnh đó, sau mỗi học kỳ, nhiều sinh viên không còn sử dụng một số tài liệu hoặc thiết bị nhưng chưa có kênh thanh lý phù hợp.")
    p(doc, "Hoạt động mua bán hiện tại chủ yếu diễn ra trên nhóm Facebook, Zalo hoặc trao đổi trực tiếp. Cách làm này tồn tại nhiều hạn chế như khó tìm kiếm, khó kiểm tra trạng thái sản phẩm, khó quản lý đơn hàng, thiếu lịch sử giao dịch và thiếu công cụ thanh toán tập trung.")

    h(doc, "1.1.2. Hạn chế của các kênh mua bán truyền thống", 2)
    bullet(doc, "Thông tin sản phẩm rời rạc, dễ bị trôi bài trong các nhóm mạng xã hội.")
    bullet(doc, "Người mua khó lọc sản phẩm theo danh mục, trường học, giá hoặc tác giả.")
    bullet(doc, "Người bán khó quản lý tồn kho, đơn hàng và lịch sử giao dịch.")
    bullet(doc, "Việc thanh toán, xác nhận đơn và hỗ trợ người dùng còn thủ công.")
    bullet(doc, "Không có ví nội bộ để người dùng nạp tiền, thanh toán và rút tiền thuận tiện.")

    h(doc, "1.1.3. Giải pháp đề xuất", 2)
    p(doc, "EduShare được đề xuất như một ứng dụng mobile tập trung cho nhu cầu mua bán, trao đổi sách và dụng cụ học tập. Ứng dụng cho phép người dùng đăng bán sản phẩm, tìm kiếm, thêm vào giỏ hàng, đặt hàng, thanh toán bằng ví EduShare, chat hỗ trợ và nhận thông báo. Hệ thống admin hỗ trợ quản lý người dùng, sản phẩm, đơn hàng và các yêu cầu ví.")

    h(doc, "1.2. Mục tiêu của đề tài", 1)
    h(doc, "1.2.1. Mục tiêu tổng quát", 2)
    p(doc, "Mục tiêu tổng quát của đề tài là xây dựng một ứng dụng di động đa nền tảng hỗ trợ sinh viên mua bán sách và dụng cụ học tập, có khả năng quản lý sản phẩm, giỏ hàng, đơn hàng, ví nội bộ và thanh toán nạp ví qua PayOS.")
    h(doc, "1.2.2. Mục tiêu cụ thể", 2)
    for item in [
        "Xây dựng giao diện mobile bằng Flutter, phù hợp thao tác trên điện thoại.",
        "Tích hợp Firebase Authentication để đăng nhập và quản lý người dùng.",
        "Thiết kế dữ liệu trên Cloud Firestore cho users, products, orders, walletRequests, conversations, notifications.",
        "Xây dựng giỏ hàng và luồng đặt hàng với kiểm tra tồn kho.",
        "Xây dựng ví EduShare cho phép nạp tiền, rút tiền và thanh toán đơn hàng.",
        "Tích hợp PayOS để tạo mã thanh toán và kiểm tra trạng thái nạp ví.",
        "Xây dựng trang admin để quản lý người dùng, sản phẩm, đơn hàng và yêu cầu ví.",
        "Bổ sung hiệu ứng giao diện như chuyển tab, xác nhận thanh toán thành công nhằm tăng trải nghiệm người dùng.",
    ]:
        bullet(doc, item)

    h(doc, "1.3. Ý nghĩa thực tiễn", 1)
    h(doc, "1.3.1. Đối với người sử dụng", 3)
    p(doc, "EduShare giúp sinh viên tiết kiệm chi phí học tập thông qua việc mua lại sản phẩm đã qua sử dụng với giá phù hợp. Ứng dụng cũng giúp người bán dễ dàng thanh lý tài liệu, dụng cụ học tập không còn sử dụng, góp phần giảm lãng phí và tăng khả năng tái sử dụng tài nguyên học tập.")
    h(doc, "1.3.2. Đối với bản thân sinh viên thực hiện", 3)
    p(doc, "Quá trình xây dựng EduShare giúp sinh viên thực hiện củng cố kiến thức về phân tích yêu cầu, thiết kế hệ thống, xây dựng ứng dụng Flutter, tích hợp Firebase, tổ chức dữ liệu NoSQL, quản lý trạng thái, xử lý thanh toán và thiết kế UI/UX.")

    h(doc, "1.4. Mô tả hệ thống", 1)
    p(doc, "Hệ thống EduShare được thiết kế theo mô hình client-cloud. Ứng dụng Flutter đóng vai trò client, xử lý giao diện và tương tác người dùng. Firebase đóng vai trò backend-as-a-service, cung cấp xác thực, lưu trữ dữ liệu, thông báo và dữ liệu thời gian thực. PayOS được tích hợp để phục vụ chức năng nạp tiền vào ví EduShare.")
    for item in [
        "Phân hệ tài khoản và hồ sơ: đăng nhập, tạo hồ sơ, cập nhật thông tin cá nhân, địa chỉ giao hàng, tài khoản ngân hàng.",
        "Phân hệ sản phẩm: đăng sản phẩm, xem danh sách, tìm kiếm, yêu thích, quản lý tồn kho.",
        "Phân hệ giỏ hàng và đơn hàng: thêm sản phẩm, điều chỉnh số lượng, đặt hàng, theo dõi trạng thái.",
        "Phân hệ ví EduShare: nạp tiền, rút tiền, thanh toán đơn hàng, xác nhận thanh toán.",
        "Phân hệ chat và thông báo: hỗ trợ người dùng, nhắn tin với admin, thông báo trạng thái.",
        "Phân hệ admin: quản lý người dùng, sản phẩm, đơn hàng, yêu cầu ví và giải ngân.",
    ]:
        bullet(doc, item)

    h(doc, "1.5. Công nghệ sử dụng", 1)
    h(doc, "1.5.1. Môi trường và ngôn ngữ lập trình", 3)
    bullet(doc, "Môi trường phát triển: Visual Studio Code, Android Emulator, Git/GitHub.")
    bullet(doc, "Ngôn ngữ lập trình: Dart.")
    bullet(doc, "Framework: Flutter.")
    h(doc, "1.5.2. Kiến trúc phần mềm", 3)
    bullet(doc, "Component-based UI: giao diện chia thành các màn hình và widget độc lập.")
    bullet(doc, "Service Layer: FirebaseDataService, PayosService, NotificationSystemService xử lý nghiệp vụ.")
    bullet(doc, "Provider: quản lý trạng thái đăng nhập và giỏ hàng.")
    h(doc, "1.5.3. Công nghệ cơ sở dữ liệu và đám mây", 3)
    bullet(doc, "Firebase Authentication: xác thực người dùng.")
    bullet(doc, "Cloud Firestore: lưu trữ users, products, orders, walletRequests, conversations, notifications.")
    bullet(doc, "Firebase Messaging/Local Notifications: hỗ trợ thông báo.")
    h(doc, "1.5.4. Tích hợp thanh toán", 3)
    bullet(doc, "PayOS API: tạo payment link, QR thanh toán và kiểm tra trạng thái giao dịch.")
    bullet(doc, "HMAC SHA256: ký dữ liệu gửi sang PayOS.")
    h(doc, "1.5.5. Thư viện giao diện", 3)
    bullet(doc, "Material Design Components: AppBar, Card, BottomAppBar, FloatingActionButton, BottomSheet.")
    bullet(doc, "CustomPainter: tạo hiệu ứng thanh toán thành công.")
    bullet(doc, "Flutter Map: chọn vị trí giao hàng.")


def chapter2(doc):
    h(doc, "CHƯƠNG 2 CƠ SỞ LÝ THUYẾT", 1)
    sections = [
        ("2.1. Framework Flutter và ngôn ngữ lập trình Dart", "Flutter là bộ SDK phát triển ứng dụng đa nền tảng do Google phát triển. Flutter sử dụng Dart và cơ chế widget tree để xây dựng giao diện. Ưu điểm của Flutter là khả năng hot reload, hiệu năng tốt, giao diện đồng nhất trên nhiều nền tảng và hệ sinh thái thư viện phong phú."),
        ("2.2. Firebase Backend-as-a-Service", "Firebase cung cấp nhiều dịch vụ backend như Authentication, Cloud Firestore, Cloud Messaging. Trong EduShare, Firebase giúp rút ngắn thời gian xây dựng backend, phù hợp với đồ án mobile cần triển khai nhanh nhưng vẫn có dữ liệu thời gian thực."),
        ("2.3. Cloud Firestore và mô hình NoSQL", "Firestore tổ chức dữ liệu theo collection và document. Mô hình này phù hợp với ứng dụng có nhiều đối tượng linh hoạt như sản phẩm, đơn hàng, yêu cầu ví, tin nhắn và thông báo. Khác với cơ sở dữ liệu quan hệ, Firestore không bắt buộc khóa ngoại nhưng cần thiết kế ID liên kết logic giữa các document."),
        ("2.4. Quản lý trạng thái với Provider", "Provider là thư viện quản lý trạng thái phổ biến trong Flutter. EduShare dùng Provider để quản lý AuthProvider và CartProvider, giúp giao diện tự cập nhật khi trạng thái đăng nhập hoặc giỏ hàng thay đổi."),
        ("2.5. Thanh toán điện tử và PayOS", "PayOS hỗ trợ tạo link thanh toán, mã QR và kiểm tra trạng thái giao dịch. Trong EduShare, PayOS được dùng cho chức năng nạp ví. Khi PayOS trả về trạng thái PAID, hệ thống cập nhật yêu cầu nạp tiền và cộng số dư ví cho người dùng."),
        ("2.6. Bảo mật dữ liệu trong ứng dụng mobile", "Ứng dụng mobile cần kiểm soát quyền đọc/ghi dữ liệu bằng Firebase Security Rules. Các thao tác liên quan ví và đơn hàng cần được kiểm tra theo request.auth.uid. Trong môi trường production, các khóa thanh toán nên được đặt ở backend thay vì client."),
        ("2.7. UML và mô hình hóa hệ thống", "UML cung cấp các sơ đồ như Use-case, Class Diagram, Activity Diagram, State Machine Diagram, Communication Diagram và Sequence Diagram. Các sơ đồ này giúp mô tả rõ yêu cầu, cấu trúc và hành vi hệ thống trước khi triển khai."),
        ("2.8. Mô hình C4", "C4 giúp mô tả kiến trúc hệ thống qua nhiều cấp độ: Context, Container, Component và Code. Trong báo cáo này, C4 được dùng để mô tả cách ứng dụng Flutter giao tiếp với Firebase, PayOS và các thành phần nội bộ."),
    ]
    for title, body in sections:
        h(doc, title, 1)
        p(doc, body)


def chapter3(doc):
    h(doc, "CHƯƠNG 3 PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    h(doc, "3.1 Phân tích yêu cầu hệ thống", 1)
    h(doc, "3.1.1. Yêu cầu chức năng (Functional Requirements)", 3)
    table(doc, ["Phân hệ", "Yêu cầu chức năng"], [
        ["Tài khoản và hồ sơ", "Đăng nhập, tạo hồ sơ, cập nhật thông tin cá nhân, địa chỉ giao hàng, tài khoản ngân hàng."],
        ["Sản phẩm", "Đăng sản phẩm, xem danh sách, tìm kiếm, yêu thích, quản lý tồn kho."],
        ["Giỏ hàng và đơn hàng", "Thêm/xóa sản phẩm, tăng giảm số lượng, tạo đơn, cập nhật trạng thái."],
        ["Ví EduShare", "Nạp tiền, rút tiền, thanh toán bằng ví, tự động cộng ví sau khi PayOS xác nhận."],
        ["Chat và thông báo", "Chat admin, lưu tin nhắn, tạo thông báo trạng thái."],
        ["Admin", "Quản lý user, sản phẩm, đơn hàng, yêu cầu ví, giải ngân."],
    ])
    h(doc, "3.1.2. Yêu cầu phi chức năng (Non-functional Requirements)", 3)
    bullet(doc, "Hiệu năng: thao tác chuyển tab, thêm giỏ hàng, tìm kiếm và mở màn hình cần phản hồi nhanh.")
    bullet(doc, "Bảo mật: dữ liệu người dùng và ví phải được giới hạn bằng Firebase Auth và Firestore Rules.")
    bullet(doc, "Tính toàn vẹn: thanh toán ví và cộng tiền nạp phải dùng transaction hoặc FieldValue.increment.")
    bullet(doc, "Tính khả dụng: giao diện rõ ràng, dễ thao tác bằng một tay, có phản hồi khi thành công/thất bại.")
    bullet(doc, "Khả năng mở rộng: service và model tách riêng để dễ bổ sung backend, webhook và chức năng mới.")

    h(doc, "3.2 Sơ đồ Use-case (Use-case Diagram)", 1)
    p(doc, "Sơ đồ Use-case tổng quan mô tả các tương tác giữa Người dùng, Người bán, Admin và các dịch vụ ngoài như Firebase, PayOS.")
    image_note(doc, "Vẽ sơ đồ Use-case tổng quan gồm actor Người dùng, Người bán, Admin, Firebase, PayOS.")
    table(doc, ["Actor", "Use-case chính"], [
        ["Người dùng", "Đăng nhập, xem/tìm kiếm sản phẩm, yêu thích, thêm giỏ hàng, đặt hàng, thanh toán, nạp/rút ví, chat."],
        ["Người bán", "Đăng sản phẩm, quản lý sản phẩm, nhận giải ngân."],
        ["Admin", "Quản lý người dùng, sản phẩm, đơn hàng, ví và thông báo."],
        ["PayOS", "Tạo link thanh toán và trả trạng thái giao dịch."],
        ["Firebase", "Xác thực, lưu trữ dữ liệu, đồng bộ realtime."],
    ])

    h(doc, "3.3. Kiến trúc hệ thống theo mô hình C4", 1)
    h(doc, "3.3.1. Cấp độ 1: Sơ đồ Ngữ cảnh (System Context Diagram)", 3)
    p(doc, "Ở mức ngữ cảnh, EduShare là ứng dụng mobile cho sinh viên. Người dùng tương tác với ứng dụng Flutter, ứng dụng giao tiếp với Firebase để xác thực/lưu dữ liệu và PayOS để xử lý nạp ví.")
    code(doc, "User/Admin -> EduShare Mobile App\nEduShare Mobile App -> Firebase Auth/Firestore\nEduShare Mobile App -> PayOS API\nEduShare Mobile App -> Firebase Messaging")
    image_note(doc, "Chèn sơ đồ Context Diagram EduShare - User/Admin - Firebase - PayOS.")
    h(doc, "3.3.2. Sơ đồ Container (Container Diagram)", 3)
    table(doc, ["Container", "Mô tả"], [
        ["Flutter Mobile App", "Giao diện, xử lý tương tác, gọi service."],
        ["Firebase Authentication", "Xác thực người dùng."],
        ["Cloud Firestore", "Lưu dữ liệu users, products, orders, walletRequests, chat, notifications."],
        ["PayOS API", "Tạo link thanh toán và kiểm tra trạng thái."],
        ["Firebase Messaging", "Hỗ trợ thông báo."],
    ])
    image_note(doc, "Chèn Container Diagram thể hiện Flutter App, Firebase, PayOS.")
    h(doc, "3.3.3. Cấp độ 3: Sơ đồ Thành phần (Component Diagram)", 3)
    table(doc, ["Component", "File/Thư mục", "Vai trò"], [
        ["Screens", "lib/screens", "Giao diện các màn hình."],
        ["Providers", "lib/providers", "Quản lý trạng thái auth và cart."],
        ["Models", "lib/models", "Định nghĩa dữ liệu."],
        ["FirebaseDataService", "lib/services/firebase_data_service.dart", "Xử lý nghiệp vụ Firestore."],
        ["PayosService", "lib/services/payos_service.dart", "Tích hợp PayOS."],
        ["Utils", "lib/utils", "Hằng số, helper, định dạng."],
    ])
    image_note(doc, "Chèn Component Diagram các thành phần Flutter nội bộ.")

    h(doc, "3.4. Thiết kế chi tiết cấu trúc hệ thống lớp (Class Diagram)", 1)
    table(doc, ["Lớp/Model", "Thuộc tính chính", "Vai trò"], [
        ["UserProfile", "id, name, email, walletBalance, isAdmin, shippingAddress", "Thông tin người dùng và ví."],
        ["Product", "id, title, price, stockQuantity, sellerUid, imageUrl", "Sản phẩm được đăng bán."],
        ["CartItem", "product, quantity, totalPrice", "Sản phẩm trong giỏ hàng."],
        ["PurchaseRecord", "buyerUid, sellerUid, productId, totalPrice, status", "Đơn hàng/lịch sử mua."],
        ["WalletRequest", "userUid, type, requestedAmount, creditedAmount, status, payosStatus", "Yêu cầu nạp/rút ví."],
        ["ChatConversation", "participantIds, lastMessage, updatedAt", "Cuộc trò chuyện."],
        ["ChatMessage", "senderUid, text, createdAt", "Tin nhắn."],
        ["AppNotification", "userUid, title, body, type, isRead", "Thông báo."],
    ])
    image_note(doc, "Chèn Class Diagram các model và service chính.")

    h(doc, "3.5. Thiết kế các sơ đồ hành vi", 1)
    h(doc, "3.5.1. Sơ đồ Hoạt động (Activity Diagram)", 3)
    code(doc, "Mở app -> Đăng nhập -> Xem sản phẩm -> Thêm giỏ hàng -> Chọn thanh toán -> Kiểm tra điều kiện -> Tạo đơn -> Cập nhật tồn kho -> Hiển thị kết quả")
    image_note(doc, "Chèn Activity Diagram luồng đặt hàng và thanh toán bằng ví.")
    h(doc, "3.5.2. Sơ đồ Trạng thái (State Machine Diagram)", 3)
    table(doc, ["Đối tượng", "Trạng thái"], [
        ["WalletRequest", "pending -> completed/cancelled"],
        ["Order", "pending_cod -> awaiting_shipment -> delivered_pending_release -> completed"],
        ["Notification", "unread -> read"],
        ["Product", "available -> out_of_stock"],
    ])
    image_note(doc, "Chèn State Machine Diagram cho Order và WalletRequest.")

    h(doc, "3.6. Sơ đồ Giao tiếp (Communication Diagram)", 1)
    code(doc, "CartScreen -> FirebaseDataService -> Firestore.orders\nCartScreen -> CartProvider\nProfileScreen -> FirebaseDataService -> PayosService -> PayOS API\nAdminDashboardScreen -> FirebaseDataService -> Firestore")
    image_note(doc, "Chèn Communication Diagram giữa Screens, Services, Firestore, PayOS.")

    h(doc, "3.7. Sơ đồ tuần tự chi tiết (Sequence Diagram)", 1)
    p(doc, "Các sơ đồ tuần tự dưới đây mô tả luồng xử lý các chức năng trọng tâm của EduShare. Có thể vẽ lại bằng PlantUML/draw.io để chèn hình chính thức.")
    h(doc, "3.7.1. Sequence Diagram: Nạp ví qua PayOS", 3)
    code(doc, "User -> ProfileScreen: Chọn Nạp tiền\nProfileScreen -> FirebaseDataService: requestWalletDeposit(amount)\nFirebaseDataService -> Firestore.walletRequests: create pending\nFirebaseDataService -> PayosService: createWalletTopupLink()\nPayosService -> PayOS API: POST payment-requests\nPayOS API --> PayosService: qrCode, checkoutUrl, paymentLinkId\nFirebaseDataService -> Firestore.walletRequests: update PayOS data\nWalletTopupScreen -> FirebaseDataService: autoConfirmWalletDepositFromBankTransaction(id)\nFirebaseDataService -> PayosService: getPaymentLink(id)\nPayosService -> PayOS API: GET payment status\nalt status == PAID\n  FirebaseDataService -> Firestore transaction: completed + increment walletBalance\n  WalletTopupScreen -> User: Success effect\nelse unpaid\n  WalletTopupScreen -> User: Continue checking\nend")
    image_note(doc, "Chèn Sequence Diagram Nạp ví PayOS.")
    h(doc, "3.7.2. Sequence Diagram: Thanh toán bằng ví EduShare", 3)
    code(doc, "User -> CartScreen: Chọn thanh toán\nCartScreen -> FirebaseDataService: getCurrentUserProfile()\nCartScreen -> FirebaseDataService: createWalletPaidOrdersFromCart(items)\nFirebaseDataService -> Firestore.products: load stock\nalt walletBalance đủ và stock đủ\n  FirebaseDataService -> Firestore.orders: create orders\n  FirebaseDataService -> Firestore.users: subtract walletBalance\n  FirebaseDataService -> Firestore.products: update stock\n  CartScreen -> CartProvider: clearCart()\n  CartScreen -> User: Success dialog\nelse invalid\n  CartScreen -> User: Show error\nend")
    image_note(doc, "Chèn Sequence Diagram Thanh toán bằng ví.")
    h(doc, "3.7.3. Sequence Diagram: Chat hỗ trợ admin", 3)
    code(doc, "User -> ProfileScreen: Chat admin\nProfileScreen -> FirebaseDataService: ensureAdminConversation()\nFirebaseDataService -> Firestore.conversations: create/get conversation\nUser -> ChatScreen: send message\nChatScreen -> FirebaseDataService: sendChatMessage()\nFirebaseDataService -> Firestore.messages: add message\nFirebaseDataService -> Firestore.notifications: create notification\nFirebaseDataService -> SupportBotService: replyFor(text)\nSupportBotService --> FirebaseDataService: bot reply\nFirebaseDataService -> Firestore.messages: add bot message")
    image_note(doc, "Chèn Sequence Diagram Chat hỗ trợ admin.")

    h(doc, "3.8. Thiết kế cấu trúc cơ sở dữ liệu chi tiết", 1)
    table(doc, ["Collection", "Trường tiêu biểu", "Mô tả"], [
        ["users", "name, email, phone, walletBalance, isAdmin, shippingAddress", "Hồ sơ người dùng."],
        ["products", "title, category, price, stockQuantity, sellerUid, imageUrl", "Sản phẩm đăng bán."],
        ["orders", "buyerUid, sellerUid, productId, quantity, totalPrice, status", "Đơn hàng."],
        ["walletRequests", "userUid, type, requestedAmount, creditedAmount, status, payosOrderCode", "Yêu cầu ví."],
        ["favorites", "userUid, productId", "Sản phẩm yêu thích."],
        ["conversations/messages", "participantIds, senderUid, text, createdAt", "Chat."],
        ["notifications", "userUid, title, body, type, isRead", "Thông báo."],
    ])
    image_note(doc, "Chèn ERD logic Firestore: users-products-orders-walletRequests-chat-notifications.")


def chapter4(doc):
    h(doc, "CHƯƠNG 4 TRIỂN KHAI VÀ XÂY DỰNG ỨNG DỤNG", 1)
    h(doc, "4.1. Môi trường phát triển và Tổ chức mã nguồn", 1)
    h(doc, "4.1.1. Môi trường và công cụ phát triển", 3)
    bullet(doc, "Visual Studio Code: soạn thảo, debug Flutter.")
    bullet(doc, "Android Emulator: chạy thử ứng dụng mobile.")
    bullet(doc, "Firebase Console: quản lý Authentication, Firestore, Rules.")
    bullet(doc, "PayOS Dashboard: lấy Client ID, API Key, Checksum Key và kiểm tra giao dịch.")
    bullet(doc, "Git/GitHub: quản lý phiên bản mã nguồn.")
    h(doc, "4.1.2. Tổ chức cấu trúc thư mục", 3)
    table(doc, ["Thư mục", "Nội dung"], [
        ["lib/screens", "Các màn hình: Home, Search, Cart, Profile, Admin, Chat."],
        ["lib/services", "FirebaseDataService, PayosService, NotificationSystemService, SupportBotService."],
        ["lib/models", "Product, UserProfile, WalletRequest, PurchaseRecord, ChatMessage..."],
        ["lib/providers", "AuthProvider, CartProvider."],
        ["lib/utils", "Constants, helpers, formatter."],
        ["assets/images", "Ảnh tĩnh của ứng dụng."],
    ])

    h(doc, "4.2. Cấu trúc thư mục dự án (Project Structure)", 1)
    h(doc, "4.2.1. Khởi tạo và cấu hình Firebase", 3)
    p(doc, "Ứng dụng khởi tạo Firebase ở main.dart và sử dụng FirebaseAuth/CloudFirestore trong FirebaseDataService. Người dùng sau khi đăng nhập sẽ được đảm bảo có hồ sơ thông qua hàm ensureUserProfile.")
    h(doc, "4.2.2. Thiết lập PayOS", 3)
    p(doc, "PayOS được cấu hình bằng dart-define thông qua PAYOS_CLIENT_ID, PAYOS_API_KEY và PAYOS_CHECKSUM_KEY. PayosService tạo chữ ký HMAC SHA256 và gọi API tạo payment link.")
    h(doc, "4.2.3. Quản lý trạng thái", 3)
    p(doc, "AuthProvider quản lý trạng thái đăng nhập. CartProvider quản lý danh sách CartItem, tổng số lượng và tổng tiền. AppShell lắng nghe CartProvider để hiển thị badge giỏ hàng.")

    h(doc, "4.3. Triển khai Lớp Nghiệp vụ (Domain/Service Layer)", 1)
    h(doc, "4.3.1. FirebaseDataService", 3)
    p(doc, "FirebaseDataService là lớp nghiệp vụ trung tâm, cung cấp các hàm đọc/ghi dữ liệu cho sản phẩm, đơn hàng, ví, chat, thông báo và admin.")
    bullet(doc, "requestWalletDeposit: tạo yêu cầu nạp ví và tạo link PayOS.")
    bullet(doc, "autoConfirmWalletDepositFromBankTransaction: kiểm tra trạng thái PayOS hoặc giao dịch ngân hàng.")
    bullet(doc, "_completeWalletDeposit: dùng transaction để đổi trạng thái request và cộng ví.")
    bullet(doc, "createWalletPaidOrdersFromCart: tạo đơn thanh toán bằng ví và cập nhật tồn kho.")
    h(doc, "4.3.2. PayosService", 3)
    p(doc, "PayosService chịu trách nhiệm giao tiếp PayOS. Lớp này tạo payment link, kiểm tra trạng thái payment link và log lỗi khi PayOS trả về lỗi.")
    h(doc, "4.3.3. Notification và Chat", 3)
    p(doc, "Ứng dụng tạo thông báo cho các sự kiện như nạp ví hoàn tất, đơn hàng đã thanh toán, tin nhắn mới. Chat lưu dữ liệu trong conversations và subcollection messages.")

    h(doc, "4.4. Triển khai Lớp Hiển thị và Quản lý Trạng thái", 1)
    h(doc, "4.4.1. AppShell và hiệu ứng chuyển tab", 3)
    p(doc, "AppShell quản lý 4 tab chính: Trang chủ, Tìm kiếm, Giỏ hàng và Hồ sơ. Hàm _buildAnimatedTabBody dùng AnimatedOpacity, AnimatedSlide và AnimatedScale để tạo hiệu ứng chuyển tab.")
    h(doc, "4.4.2. ProfileScreen và ví EduShare", 3)
    p(doc, "ProfileScreen hiển thị hồ sơ, ví, nạp tiền, rút tiền, chat admin. Hiệu ứng thanh toán thành công được triển khai bằng CustomPainter và overlay.")
    h(doc, "4.4.3. CartScreen và thanh toán", 3)
    p(doc, "CartScreen hiển thị sản phẩm trong giỏ, chọn phương thức thanh toán COD hoặc ví EduShare. Chức năng chuyển khoản QR tự động đã được gỡ bỏ để tránh nhầm lẫn khi chưa có backend đối soát ngân hàng.")

    h(doc, "4.5. Giao diện Ứng dụng thực tế (Chương trình Demo)", 1)
    h(doc, "4.5.1. Màn hình Trang chủ", 3)
    p(doc, "Trang chủ hiển thị sản phẩm mới, sản phẩm nổi bật và các gợi ý phù hợp với người dùng.")
    image_note(doc, "Chèn ảnh màn hình Trang chủ có danh sách sản phẩm.")
    h(doc, "4.5.2. Màn hình Chi tiết sản phẩm", 3)
    p(doc, "Màn hình chi tiết sản phẩm hiển thị ảnh, tên, giá, tồn kho, mô tả và thông tin người bán.")
    image_note(doc, "Chèn ảnh chi tiết sản phẩm có ảnh sản phẩm rõ ràng.")
    h(doc, "4.5.3. Màn hình Giỏ hàng và thanh toán", 3)
    p(doc, "Giỏ hàng cho phép điều chỉnh số lượng và chọn phương thức thanh toán. Sau chỉnh sửa, chỉ còn COD và ví EduShare.")
    image_note(doc, "Chèn ảnh giỏ hàng và bottom sheet chọn thanh toán.")
    h(doc, "4.5.4. Màn hình Ví EduShare và PayOS", 3)
    p(doc, "Màn hình ví hiển thị số dư, nạp tiền, rút tiền. Khi nạp tiền, ứng dụng tạo QR PayOS và tự kiểm tra trạng thái.")
    image_note(doc, "Chèn ảnh ví EduShare, màn hình QR PayOS và hiệu ứng thanh toán thành công.")
    h(doc, "4.5.5. Màn hình Admin Dashboard", 3)
    p(doc, "Admin Dashboard hỗ trợ quản lý người dùng, sản phẩm, đơn hàng và yêu cầu ví.")
    image_note(doc, "Chèn ảnh Admin Dashboard phần đơn hàng hoặc yêu cầu ví.")
    h(doc, "4.5.6. Các vị trí cần thêm hình ảnh sản phẩm", 3)
    table(doc, ["Vị trí", "Ảnh cần thêm"], [
        ["Trang chủ", "Ảnh danh sách sản phẩm: sách, giáo trình, dụng cụ học tập."],
        ["Chi tiết sản phẩm", "Ảnh một sản phẩm cụ thể với giá và tồn kho."],
        ["Đăng sản phẩm", "Ảnh form đăng sản phẩm có upload/chọn ảnh."],
        ["Giỏ hàng", "Ảnh sản phẩm đã thêm vào giỏ."],
        ["Báo cáo dữ liệu products", "Ảnh minh họa các loại sản phẩm: sách, máy tính, dụng cụ."],
    ])


def chapter5(doc):
    h(doc, "CHƯƠNG 5 KIỂM THỬ VÀ ĐÁNH GIÁ HỆ THỐNG", 1)
    h(doc, "5.1. Kiểm thử hệ thống (System Testing)", 1)
    h(doc, "5.1.1. Mục đích và phương pháp kiểm thử", 3)
    p(doc, "Kiểm thử nhằm đảm bảo các chức năng chính hoạt động đúng, dữ liệu được lưu vào Firestore, ví cập nhật chính xác và giao diện phản hồi đúng với thao tác người dùng. Phương pháp kiểm thử chủ yếu là kiểm thử thủ công trên emulator kết hợp theo dõi dữ liệu trên Firebase Console.")
    h(doc, "5.1.2. Các kịch bản kiểm thử cốt lõi", 3)
    table(doc, ["Mã TC", "Chức năng", "Thao tác", "Kết quả mong đợi"], [
        ["TC01", "Đăng nhập", "Nhập tài khoản hợp lệ", "Vào được AppShell và tạo hồ sơ nếu chưa có."],
        ["TC02", "Đăng sản phẩm", "Nhập thông tin sản phẩm", "Sản phẩm được lưu vào products."],
        ["TC03", "Tìm kiếm", "Nhập từ khóa sản phẩm", "Danh sách kết quả phù hợp."],
        ["TC04", "Giỏ hàng", "Thêm/tăng/giảm số lượng", "Tổng tiền và badge cập nhật đúng."],
        ["TC05", "Thanh toán ví đủ tiền", "Chọn ví EduShare", "Tạo đơn, trừ ví, cập nhật tồn kho."],
        ["TC06", "Thanh toán ví thiếu tiền", "Số dư nhỏ hơn tổng tiền", "Hiển thị lỗi số dư không đủ."],
        ["TC07", "Nạp ví PayOS", "Tạo yêu cầu nạp", "Tạo walletRequest và QR PayOS."],
        ["TC08", "PayOS PAID", "Giao dịch hoàn tất", "walletRequest completed, walletBalance tăng."],
        ["TC09", "Rút tiền", "Tạo yêu cầu rút", "walletRequest withdrawal pending, ví tạm trừ."],
        ["TC10", "Chat admin", "Gửi tin nhắn", "Tin nhắn lưu realtime và tạo thông báo."],
    ])

    h(doc, "5.2. Kết quả đạt được", 1)
    bullet(doc, "Hoàn thành ứng dụng Flutter có luồng đăng nhập, hồ sơ, sản phẩm, giỏ hàng, ví và admin.")
    bullet(doc, "Tích hợp Firebase Authentication và Cloud Firestore.")
    bullet(doc, "Tích hợp PayOS cho nạp tiền vào ví.")
    bullet(doc, "Xây dựng hiệu ứng chuyển tab và hiệu ứng thanh toán thành công.")
    bullet(doc, "Gỡ bỏ phương thức chuyển khoản QR tự động trong giỏ hàng để tránh nhầm lẫn khi chưa có backend đối soát.")

    h(doc, "5.3. Đánh giá hệ thống", 1)
    h(doc, "5.3.1. So sánh kết quả thực tế với mục tiêu ban đầu", 3)
    p(doc, "Ứng dụng đã đáp ứng phần lớn mục tiêu ban đầu: người dùng có thể đăng nhập, xem sản phẩm, đăng bán, đặt hàng, thanh toán bằng ví, nạp ví qua PayOS, chat và nhận thông báo. Admin có thể quản lý dữ liệu trọng tâm.")
    h(doc, "5.3.2. Ưu điểm nổi bật của hệ thống", 3)
    bullet(doc, "Tổ chức code rõ theo screens, services, models, providers.")
    bullet(doc, "Firebase giúp đồng bộ dữ liệu nhanh và giảm khối lượng backend cần tự viết.")
    bullet(doc, "Ví EduShare và PayOS tạo điểm khác biệt so với ứng dụng đăng bán đơn giản.")
    bullet(doc, "Giao diện có phản hồi trực quan, hiệu ứng chuyển tab và thanh toán thành công.")
    h(doc, "5.3.3. Hạn chế còn tồn tại", 3)
    bullet(doc, "PayOS hiện vẫn được gọi từ client để phục vụ demo; production nên đưa sang backend/cloud function.")
    bullet(doc, "Chưa có đánh giá người bán, báo cáo vi phạm và lọc nâng cao.")
    bullet(doc, "Chưa có kiểm thử tự động đầy đủ.")
    bullet(doc, "Một số dữ liệu ảnh sản phẩm cần bổ sung thêm để báo cáo và demo trực quan hơn.")


def chapter6(doc):
    h(doc, "CHƯƠNG 6 KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    h(doc, "6.1. Tóm tắt nội dung chính của đề tài", 1)
    p(doc, "Đề tài đã xây dựng ứng dụng mobile EduShare phục vụ nhu cầu mua bán sách và dụng cụ học tập của sinh viên. Hệ thống gồm các phân hệ tài khoản, sản phẩm, giỏ hàng, đơn hàng, ví EduShare, nạp ví qua PayOS, chat hỗ trợ, thông báo và admin dashboard.")
    p(doc, "Thông qua quá trình thực hiện, sinh viên đã vận dụng được kiến thức về Flutter, Firebase, Firestore, Provider, tích hợp API thanh toán, thiết kế dữ liệu NoSQL và thiết kế giao diện người dùng.")
    h(doc, "6.2. Định hướng mở rộng trong tương lai", 1)
    h(doc, "6.2.1. Chuyển xử lý thanh toán sang backend", 3)
    p(doc, "Triển khai Firebase Cloud Functions hoặc backend riêng để tạo payment link, xác thực webhook PayOS và cộng tiền ví, tránh để API key ở client.")
    h(doc, "6.2.2. Bổ sung đánh giá và độ uy tín", 3)
    p(doc, "Thêm chức năng đánh giá người bán/người mua, báo cáo sản phẩm vi phạm và thống kê độ uy tín.")
    h(doc, "6.2.3. Nâng cấp tìm kiếm và gợi ý", 3)
    p(doc, "Bổ sung lọc theo khoảng giá, trường học, loại sản phẩm, tình trạng sản phẩm và gợi ý thông minh dựa trên hành vi người dùng.")
    h(doc, "6.2.4. Mở rộng hệ sinh thái", 3)
    p(doc, "Phát triển phiên bản web admin, thống kê doanh thu, quản lý vận chuyển và tối ưu giao diện cho nhiều kích thước màn hình.")


def refs(doc):
    h(doc, "TÀI LIỆU THAM KHẢO", 1)
    for ref in [
        "Flutter Documentation: https://docs.flutter.dev/",
        "Dart Documentation: https://dart.dev/guides",
        "Firebase Documentation: https://firebase.google.com/docs",
        "Cloud Firestore Documentation: https://firebase.google.com/docs/firestore",
        "PayOS Documentation: https://payos.vn/docs/",
        "Provider Package: https://pub.dev/packages/provider",
        "Flutter Map Package: https://pub.dev/packages/flutter_map",
    ]:
        p(doc, ref)


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    preface(doc)
    lists(doc)
    chapter1(doc)
    page(doc)
    chapter2(doc)
    page(doc)
    chapter3(doc)
    page(doc)
    chapter4(doc)
    page(doc)
    chapter5(doc)
    page(doc)
    chapter6(doc)
    page(doc)
    refs(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
